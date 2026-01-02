from __future__ import annotations

import asyncio
import json
import logging
import os
import sys
import time
from abc import ABC, abstractmethod
from collections import deque
from concurrent.futures import ThreadPoolExecutor
from contextlib import asynccontextmanager, contextmanager
from dataclasses import dataclass, field
from datetime import date
from enum import Enum, auto
from functools import lru_cache
from pathlib import Path
from typing import (
    Any,
    Callable,
    Deque,
    Dict,
    Generator,
    List,
    Optional,
    Protocol,
    Sequence,
    Tuple,
    TypeVar,
)
import argparse

# 第三方库
import google.generativeai as genai
import pymupdf
from bs4 import BeautifulSoup
from docx import Document
from ebooklib import ITEM_DOCUMENT, epub
from pathvalidate import sanitize_filename
from tqdm.asyncio import tqdm

# ============================================================================
# 配置模块
# ============================================================================

@dataclass(frozen=True)
class Config:
    """不可变的全局配置类"""
    # API 限制
    rpm_limit: int = 10
    tpm_limit: int = 250_000
    daily_request_limit: int = 250
    max_tokens_per_request: int = 27_000
    
    # 并发控制
    concurrency_limit: int = 10
    max_retries: int = 3
    max_items_per_batch: int = 12
    
    # IO 配置
    io_workers: int = field(default_factory=lambda: min(32, max(4, (os.cpu_count() or 8) * 2)))
    
    # 文本处理
    chars_per_token: float = 3.5
    
    # 文件类型
    supported_extensions: Tuple[str, ...] = ('.pdf', '.epub', '.azw3', '.docx')
    
    # 路径配置
    pending_files_log: Path = field(default_factory=lambda: Path("./pending_files.txt"))
    tracker_file: Path = field(default_factory=lambda: Path("./request_tracker.json"))
    
    @property
    def max_chars_per_request(self) -> int:
        return int(self.max_tokens_per_request * self.chars_per_token)


# 使用默认配置初始化（可通过依赖注入替换）
CONFIG = Config()


class ProcessingMode(Enum):
    """处理模式枚举"""
    BATCH = auto()
    SINGLE = auto()


# ============================================================================
# Prompts 和 Schema 定义
# ============================================================================

PROMPTS = {
    'batch': """
Analyze the following text, which contains MULTIPLE documents concatenated together.
Each document starts with a "--- START OF FILE: [filename] ---" marker and ends with an "--- END OF FILE: [filename] ---" marker.
For EACH document provided, extract its metadata. Crucially, also extract a list of 3-5 relevant keywords from the document's content.
Return a single JSON array (a list) containing all the extracted JSON objects.
The order of objects in the final list MUST match the order of the documents in the input text.
Do not add any commentary. Only return the JSON array.
""".strip(),
    
    'single': """
Analyze the text from the following document to extract its metadata.
Based on the content, provide a JSON object with the following details.
Crucially, also extract a list of 3-5 relevant keywords from the document's content.
Do not add any commentary. Only return the JSON object.
""".strip()
}

SINGLE_OBJECT_SCHEMA: Dict[str, Any] = {
    "type": "object",
    "properties": {
        "title": {"type": "string"},
        "authors": {"type": "array", "items": {"type": "string"}},
        "keywords": {"type": "array", "items": {"type": "string"}},
        "translators": {"type": "string"},
        "editors": {"type": "string"},
        "publisher_or_journal": {"type": "string"},
        "journal_volume_issue": {"type": "string"},
        "publication_date": {"type": "string"},
        "start_page": {"type": "integer"}
    },
    "required": ["title"]
}

BATCH_SCHEMA: Dict[str, Any] = {
    "type": "array",
    "items": SINGLE_OBJECT_SCHEMA
}


# ============================================================================
# 日志配置
# ============================================================================

def setup_logging(level: int = logging.INFO) -> logging.Logger:
    """配置并返回日志器"""
    logging.basicConfig(
        level=level,
        format='%(asctime)s - %(levelname)s - %(message)s',
        handlers=[logging.StreamHandler(sys.stdout)]
    )
    return logging.getLogger(__name__)


logger = setup_logging()


# ============================================================================
# 数据模型
# ============================================================================

@dataclass
class FileItem:
    """表示待处理文件的数据类"""
    path: Path
    text: str
    tokens: int
    
    def __hash__(self) -> int:
        return hash(self.path)


@dataclass
class BatchResult:
    """批处理结果"""
    success: bool
    failed_items: List[FileItem] = field(default_factory=list)
    quota_exceeded: bool = False


@dataclass
class SingleResult:
    """单文件处理结果"""
    success: bool
    failed_item: Optional[FileItem] = None
    quota_exceeded: bool = False


@dataclass
class Batch:
    """表示一个处理批次"""
    items: List[FileItem]
    tokens: int


@dataclass
class ProcessingStats:
    """处理统计信息"""
    total_processed: int = 0
    total_failed: int = 0
    prep_time: float = 0.0
    api_time: float = 0.0
    
    @property
    def total_time(self) -> float:
        return self.prep_time + self.api_time
    
    @property
    def average_rate(self) -> float:
        return self.total_processed / self.api_time if self.api_time > 0 else 0.0


# ============================================================================
# API 密钥管理
# ============================================================================

class APIKeyManager:
    """API 密钥管理器"""
    
    def __init__(self, keys: List[str], tracker_file: Path):
        self._keys = keys
        self._tracker_file = tracker_file
        self._current_index = 0
        self._tracker = self._load_tracker()
    
    def _load_tracker(self) -> Dict[str, Any]:
        """加载请求跟踪器"""
        today_str = date.today().isoformat()
        default = {"date": today_str, "usage": {}}
        
        if not self._tracker_file.exists():
            return default
            
        try:
            with open(self._tracker_file, 'r', encoding='utf-8') as f:
                tracker = json.load(f)
            
            if tracker.get("date") != today_str:
                logger.info("新的一天，重置所有API密钥的每日请求计数器。")
                return default
            
            tracker.setdefault("usage", {})
            return tracker
            
        except (json.JSONDecodeError, IOError) as e:
            logger.warning(f"读取请求跟踪文件失败: {e}")
            return default
    
    def save_tracker(self) -> None:
        """保存请求跟踪器"""
        try:
            with open(self._tracker_file, 'w', encoding='utf-8') as f:
                json.dump(self._tracker, f, indent=4, ensure_ascii=False)
        except IOError as e:
            logger.error(f"保存请求跟踪文件失败: {e}")
    
    def get_usage(self, key: str) -> int:
        """获取指定密钥的使用次数"""
        return self._tracker["usage"].get(key, 0)
    
    def increment_usage(self, key: str) -> None:
        """增加指定密钥的使用次数"""
        self._tracker["usage"][key] = self.get_usage(key) + 1
    
    @property
    def keys(self) -> List[str]:
        return self._keys
    
    @property
    def count(self) -> int:
        return len(self._keys)
    
    def get_remaining_quota(self, key: str, daily_limit: int) -> int:
        """获取剩余配额"""
        return daily_limit - self.get_usage(key)


def configure_api_keys() -> List[str]:
    """从环境变量或用户输入获取API密钥"""
    keys_str = os.getenv("GOOGLE_API_KEY")
    
    if not keys_str:
        print("-" * 65)
        print("未找到 GOOGLE_API_KEY 环境变量。")
        keys_str = input("请输入您的一个或多个 Google API 密钥 (若有多个，请用逗号','分隔):\n").strip()
        print("-" * 65)

    if not keys_str:
        logger.error("错误：未提供任何 API 密钥，程序即将退出。")
        sys.exit(1)

    api_keys = [key.strip() for key in keys_str.split(',') if key.strip()]

    if not api_keys:
        logger.error("错误：提供的 API 密钥为空，程序即将退出。")
        sys.exit(1)

    logger.info(f"找到 {len(api_keys)} 个 API 密钥。")
    return api_keys


# ============================================================================
# Gemini 模型包装器
# ============================================================================

class GeminiModel:
    """Gemini 模型包装器，封装 API 交互"""
    
    MODEL_NAME = 'models/gemini-2.5-flash'
    
    def __init__(self):
        self._model: Optional[genai.GenerativeModel] = None
        self._api_key: Optional[str] = None
    
    @property
    def api_key(self) -> Optional[str]:
        return self._api_key
    
    @property
    def is_configured(self) -> bool:
        return self._model is not None
    
    def configure(self, api_key: str) -> bool:
        """配置 API 密钥"""
        try:
            genai.configure(api_key=api_key)
            self._model = genai.GenerativeModel(self.MODEL_NAME)
            self._api_key = api_key
            logger.info(f"API 密钥 (前8位: {api_key[:8]}...) 配置成功。")
            return True
        except Exception as e:
            logger.error(f"API 密钥配置失败: {e}")
            return False
    
    def count_tokens(self, text: str) -> int:
        """计算文本的 token 数量"""
        if not self._model:
            raise RuntimeError("模型尚未配置")
        result = self._model.count_tokens(text)
        return result.total_tokens
    
    async def generate_content(
        self,
        prompt: str,
        schema: Dict[str, Any]
    ) -> str:
        """异步生成内容"""
        if not self._model:
            raise RuntimeError("模型尚未配置")
        
        config = {
            "response_mime_type": "application/json",
            "response_schema": schema
        }
        response = await self._model.generate_content_async(prompt, generation_config=config)
        return response.text


# 全局模型实例
MODEL = GeminiModel()


# ============================================================================
# 速率限制器
# ============================================================================

class RateLimiter:
    """精确速率限制器，避免轮询"""
    
    def __init__(self, rpm: int, tpm: int):
        self._rpm = rpm
        self._tpm = tpm
        self._request_timestamps: Deque[float] = deque()
        self._token_records: Deque[Tuple[float, int]] = deque()
        self._token_total = 0
        self._lock = asyncio.Lock()
    
    def _cleanup_old_records(self, now: float) -> None:
        """清理超过60秒的旧记录"""
        cutoff = now - 60
        
        while self._request_timestamps and self._request_timestamps[0] < cutoff:
            self._request_timestamps.popleft()
        
        while self._token_records and self._token_records[0][0] < cutoff:
            _, tokens = self._token_records.popleft()
            self._token_total -= tokens
    
    def _calculate_wait_time(self, now: float, tokens_needed: int) -> float:
        """计算需要等待的时间"""
        rpm_wait = 0.0
        tpm_wait = 0.0
        
        # RPM 限制检查
        if len(self._request_timestamps) >= self._rpm and self._request_timestamps:
            rpm_wait = (self._request_timestamps[0] + 60) - now
        
        # TPM 限制检查
        if (self._token_total + tokens_needed) > self._tpm and self._token_records:
            tokens_to_free = (self._token_total + tokens_needed) - self._tpm
            freed = 0
            wait_until = 0.0
            
            for ts, tk in self._token_records:
                freed += tk
                if freed >= tokens_to_free:
                    wait_until = ts
                    break
            
            if wait_until > 0:
                tpm_wait = (wait_until + 60) - now
        
        return max(0.1, rpm_wait, tpm_wait)
    
    async def acquire(self, tokens_needed: int) -> None:
        """获取 API 调用槽位"""
        async with self._lock:
            while True:
                now = time.time()
                self._cleanup_old_records(now)
                
                current_requests = len(self._request_timestamps)
                can_request = current_requests < self._rpm
                can_tokens = (self._token_total + tokens_needed) <= self._tpm
                
                if can_request and can_tokens:
                    self._request_timestamps.append(now)
                    self._token_records.append((now, tokens_needed))
                    self._token_total += tokens_needed
                    return
                
                wait_time = self._calculate_wait_time(now, tokens_needed)
                logger.info(f"速率限制，等待 {wait_time:.2f} 秒...")
                await asyncio.sleep(wait_time)


# ============================================================================
# 文本提取器
# ============================================================================

class TextExtractor(ABC):
    """文本提取器抽象基类"""
    
    @abstractmethod
    def extract(self, path: Path) -> str:
        """提取文本内容"""
        pass


class PDFExtractor(TextExtractor):
    """PDF 文本提取器"""
    
    def __init__(self, pages_start: int = 4, pages_end: int = 3):
        self._pages_start = pages_start
        self._pages_end = pages_end
    
    def extract(self, path: Path) -> str:
        try:
            with pymupdf.open(path) as doc:
                total = doc.page_count
                pages = set(range(min(self._pages_start, total)))
                
                if total > self._pages_start + self._pages_end:
                    pages.update(range(total - self._pages_end, total))
                
                texts = [doc[i].get_text(sort=True) for i in sorted(pages)]
                return "\n".join(texts)
                
        except Exception as e:
            logger.error(f"PDF 提取失败 {path.name}: {e}")
            return ""


class EPUBExtractor(TextExtractor):
    """EPUB/AZW3 文本提取器"""
    
    def __init__(self, chapters_start: int = 5, chapters_end: int = 4):
        self._chapters_start = chapters_start
        self._chapters_end = chapters_end
    
    def extract(self, path: Path) -> str:
        try:
            book = epub.read_epub(path)
            items = list(book.get_items_of_type(ITEM_DOCUMENT))
            
            to_process = items[:self._chapters_start]
            if len(items) > self._chapters_start + self._chapters_end:
                to_process.extend(items[-self._chapters_end:])
            
            texts = []
            for item in to_process:
                soup = BeautifulSoup(item.get_body_content(), 'html.parser')
                texts.append(soup.get_text("\n", strip=True))
            
            return "\n\n".join(texts)
            
        except Exception as e:
            logger.error(f"EPUB 提取失败 {path.name}: {e}")
            return ""


class DOCXExtractor(TextExtractor):
    """DOCX 文本提取器"""
    
    def __init__(self, paras_start: int = 20, paras_end: int = 15):
        self._paras_start = paras_start
        self._paras_end = paras_end
    
    def extract(self, path: Path) -> str:
        try:
            doc = Document(path)
            paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            
            if len(paras) > self._paras_start + self._paras_end:
                result = paras[:self._paras_start] + paras[-self._paras_end:]
            else:
                result = paras
            
            return "\n".join(result)
            
        except Exception as e:
            logger.error(f"DOCX 提取失败 {path.name}: {e}")
            return ""


class TextExtractorFactory:
    """文本提取器工厂"""
    
    _extractors: Dict[str, TextExtractor] = {
        '.pdf': PDFExtractor(),
        '.epub': EPUBExtractor(),
        '.azw3': EPUBExtractor(),
        '.docx': DOCXExtractor(),
    }
    
    @classmethod
    def get_extractor(cls, extension: str) -> Optional[TextExtractor]:
        return cls._extractors.get(extension.lower())


def smart_truncate(text: str, max_chars: int) -> str:
    """智能截断文本，保留头尾"""
    if len(text) <= max_chars:
        return text
    
    head_chars = int(max_chars * 0.6)
    tail_chars = int(max_chars * 0.4)
    
    return f"{text[:head_chars]}\n\n--- 内容已截断 ---\n\n{text[-tail_chars:]}"


def extract_text(path: Path, max_chars: int) -> Optional[str]:
    """从文件提取文本"""
    extractor = TextExtractorFactory.get_extractor(path.suffix)
    if not extractor:
        logger.warning(f"不支持的文件类型: {path.name}")
        return None
    
    text = extractor.extract(path)
    if not text:
        return None
    
    return smart_truncate(text, max_chars)


def extract_and_count(path: Path, config: Config) -> Optional[FileItem]:
    """提取文本并计算 token（同步函数，用于线程池）"""
    text = extract_text(path, config.max_chars_per_request)
    if not text:
        return None
    
    if not MODEL.is_configured:
        logger.error("模型尚未配置")
        return None
    
    try:
        tokens = MODEL.count_tokens(text)
        return FileItem(path=path, text=text, tokens=tokens)
    except Exception as e:
        logger.error(f"Token 计算失败 {path.name}: {e}")
        return None


# ============================================================================
# 元数据处理
# ============================================================================

# 关键词常量
JOURNAL_KEYWORDS = frozenset([
    "journal", "review", "proceedings", "transactions", "quarterly",
    "annals", "bulletin", "magazine", "advances", "letters", "studies",
    "science", "research", "technology", "medicine", "report", "archives",
    "学报", "法学", "研究", "评论", "科学", "技术", "杂志", "动态",
    "报告", "医学", "经济", "哲学", "历史", "通讯", "汇刊", "纪要"
])

UNKNOWN_AUTHOR_MARKERS = frozenset(["作者不详"])
ROLE_INVALID_TOKENS = frozenset(["null", "none", "n/a", "unknown", "不详", "未知"])
ROLE_INVALID_SUBSTRINGS = frozenset(["无法提取", "不明确", "系统返回null", "系统返回 null"])


class TextNormalizer:
    """文本规范化工具类"""
    
    @staticmethod
    def normalize(value: Any) -> str:
        """规范化单个值"""
        if value is None:
            return ""
        text = str(value).strip()
        return "" if not text or text.lower() == "null" else text
    
    @staticmethod
    def normalize_list(values: Optional[List[Any]]) -> List[str]:
        """规范化列表"""
        if not values:
            return []
        return [v for v in map(TextNormalizer.normalize, values) if v]
    
    @staticmethod
    def normalize_authors(values: Optional[List[Any]]) -> List[str]:
        """规范化作者列表"""
        authors = TextNormalizer.normalize_list(values)
        return [a for a in authors if a not in UNKNOWN_AUTHOR_MARKERS]
    
    @staticmethod
    def normalize_role(value: Any) -> str:
        """规范化角色字段（译者、编者等）"""
        normalized = TextNormalizer.normalize(value)
        if not normalized:
            return ""
        
        lower = normalized.lower()
        if lower in ROLE_INVALID_TOKENS:
            return ""
        
        compact = normalized.replace(" ", "")
        for marker in ROLE_INVALID_SUBSTRINGS:
            if marker in normalized or marker in compact:
                return ""
        
        return normalized


class MetadataBuilder:
    """元数据构建器"""
    
    def __init__(self, info: Dict[str, Any]):
        self._info = info
        self._normalizer = TextNormalizer
    
    @property
    def title(self) -> str:
        return self._normalizer.normalize(self._info.get("title"))
    
    @property
    def authors(self) -> List[str]:
        return self._normalizer.normalize_authors(self._info.get("authors"))
    
    @property
    def authors_str(self) -> str:
        return "、".join(self.authors)
    
    @property
    def keywords(self) -> List[str]:
        return self._normalizer.normalize_list(self._info.get("keywords"))
    
    @property
    def keywords_str(self) -> str:
        return ", ".join(self.keywords)
    
    @property
    def translators(self) -> str:
        return self._normalizer.normalize_role(self._info.get("translators"))
    
    @property
    def editors(self) -> str:
        return self._normalizer.normalize_role(self._info.get("editors"))
    
    @property
    def publisher(self) -> str:
        return self._normalizer.normalize(self._info.get("publisher_or_journal"))
    
    def build_details_string(self) -> str:
        """构建详细信息字符串"""
        details = []
        
        mappings = [
            ("出版/期刊", self.publisher),
            ("卷期", self._normalizer.normalize(self._info.get("journal_volume_issue"))),
            ("日期", self._normalizer.normalize(self._info.get("publication_date"))),
            ("编者", self.editors),
            ("译者", self.translators),
            ("页码", self._normalizer.normalize(self._info.get("start_page"))),
        ]
        
        for label, value in mappings:
            if value:
                details.append(f"{label}: {value}")
        
        return " | ".join(details)
    
    def build_filename(self) -> Optional[str]:
        """构建文件名
        
        引注规范：
        - 有作者时：显示作者，不显示编者（因为引注以作者为准）
        - 无作者时（集合作品）：显示编者
        - 译者始终显示（如有）
        """
        if not self.title:
            return None
        
        main_part = f"{self.title} - {self.authors_str}" if self.authors_str else self.title
        
        extras = []
        if self.translators:
            extras.append(f"{self.translators} 译")
        
        # 只有在没有作者的情况下才添加编者（集合作品/汇编作品）
        if self.editors and not self.authors:
            pub_lower = self.publisher.lower()
            # 期刊类出版物不需要编者
            if not any(k in pub_lower for k in JOURNAL_KEYWORDS):
                extras.append(f"{self.editors} 编")
        
        return f"{main_part} ({', '.join(extras)})" if extras else main_part


# ============================================================================
# 元数据写入器
# ============================================================================

class MetadataWriter(ABC):
    """元数据写入器抽象基类"""
    
    @abstractmethod
    def write(self, path: Path, builder: MetadataBuilder) -> None:
        pass


class PDFMetadataWriter(MetadataWriter):
    """PDF 元数据写入器"""
    
    def write(self, path: Path, builder: MetadataBuilder) -> None:
        try:
            with pymupdf.open(path) as doc:
                metadata = doc.metadata
                metadata['title'] = builder.title
                metadata['author'] = builder.authors_str
                metadata['subject'] = builder.build_details_string()
                metadata['keywords'] = builder.keywords_str
                
                doc.set_metadata(metadata)
                doc.save(doc.name, incremental=True, encryption=pymupdf.PDF_ENCRYPT_KEEP)
            
            logger.info(f"PDF 元数据写入成功: {path.name}")
        except Exception as e:
            logger.error(f"PDF 元数据写入失败 {path.name}: {e}")


class DOCXMetadataWriter(MetadataWriter):
    """DOCX 元数据写入器"""
    
    def write(self, path: Path, builder: MetadataBuilder) -> None:
        try:
            doc = Document(path)
            cp = doc.core_properties
            cp.title = builder.title
            cp.author = builder.authors_str
            cp.subject = builder.build_details_string()
            cp.keywords = builder.keywords_str
            cp.comments = "Metadata updated by Gemini File Renamer"
            doc.save(path)
            
            logger.info(f"DOCX 元数据写入成功: {path.name}")
        except Exception as e:
            logger.error(f"DOCX 元数据写入失败 {path.name}: {e}")


class EPUBMetadataWriter(MetadataWriter):
    """EPUB 元数据写入器"""
    
    def _clear_creators(self, book: epub.EpubBook) -> None:
        """清除现有作者信息"""
        namespace = "http://purl.org/dc/elements/1.1/"
        meta = book.metadata.get(namespace)
        
        if isinstance(meta, dict) and "creator" in meta:
            meta["creator"] = []
        elif isinstance(meta, list):
            book.metadata[namespace] = [
                item for item in meta 
                if not (isinstance(item, tuple) and item and item[0] == "creator")
            ]
    
    def write(self, path: Path, builder: MetadataBuilder) -> None:
        try:
            book = epub.read_epub(path)
            book.set_title(builder.title)
            
            self._clear_creators(book)
            for author in builder.authors:
                book.add_author(author)
            
            # 构建描述
            description_parts = []
            details = builder.build_details_string()
            if details:
                description_parts.append(details)
            if builder.keywords:
                description_parts.append(f"Keywords: {builder.keywords_str}")
            
            if description_parts:
                book.add_metadata('DC', 'description', "\n".join(description_parts))
            
            epub.write_epub(path, book)
            logger.info(f"EPUB 元数据写入成功: {path.name}")
        except Exception as e:
            logger.error(f"EPUB 元数据写入失败 {path.name}: {e}")


class MetadataWriterFactory:
    """元数据写入器工厂"""
    
    _writers: Dict[str, MetadataWriter] = {
        '.pdf': PDFMetadataWriter(),
        '.docx': DOCXMetadataWriter(),
        '.epub': EPUBMetadataWriter(),
        '.azw3': EPUBMetadataWriter(),
    }
    
    @classmethod
    def get_writer(cls, extension: str) -> Optional[MetadataWriter]:
        return cls._writers.get(extension.lower())


# ============================================================================
# 文件重命名器
# ============================================================================

class FileRenamer:
    """文件重命名和元数据写入处理器"""
    
    def __init__(self, write_metadata: bool = True):
        self._write_metadata = write_metadata
        self._executor = ThreadPoolExecutor(max_workers=4)
    
    async def process(
        self,
        path: Path,
        info: Dict[str, Any]
    ) -> None:
        """处理单个文件的重命名和元数据写入"""
        builder = MetadataBuilder(info)
        new_name = builder.build_filename()
        
        if not new_name:
            logger.warning(f"无法构建文件名: {path.name}")
            return
        
        # 清理文件名
        safe_name = sanitize_filename(new_name).strip()
        if not safe_name or safe_name in {".", ".."}:
            logger.warning(f"非法文件名: {new_name}")
            return
        
        # 构建新路径
        new_path = path.with_name(f"{safe_name}{path.suffix}")
        
        # 处理重名
        counter = 1
        while new_path.exists() and new_path != path:
            new_path = path.with_name(f"{safe_name}_{counter}{path.suffix}")
            counter += 1
        
        # 执行重命名
        if new_path != path:
            try:
                path.rename(new_path)
                logger.info(f"重命名: {path.name} -> {new_path.name}")
            except OSError as e:
                logger.error(f"重命名失败 {path.name}: {e}")
                return
        
        # 写入元数据
        if self._write_metadata:
            await self._write_metadata_async(new_path, builder)
    
    async def _write_metadata_async(
        self,
        path: Path,
        builder: MetadataBuilder
    ) -> None:
        """异步写入元数据"""
        writer = MetadataWriterFactory.get_writer(path.suffix)
        if not writer:
            return
        
        loop = asyncio.get_running_loop()
        await loop.run_in_executor(self._executor, writer.write, path, builder)


# ============================================================================
# 批处理打包器
# ============================================================================

def pack_batches_ffd(
    items: List[FileItem],
    max_tokens: int,
    max_items: Optional[int] = None
) -> List[Batch]:
    """First-Fit Decreasing 装箱算法"""
    sorted_items = sorted(items, key=lambda x: x.tokens, reverse=True)
    batches: List[Batch] = []
    
    for item in sorted_items:
        placed = False
        
        for batch in batches:
            can_fit_tokens = batch.tokens + item.tokens <= max_tokens
            can_fit_items = max_items is None or len(batch.items) < max_items
            
            if can_fit_tokens and can_fit_items:
                batch.items.append(item)
                batch.tokens += item.tokens
                placed = True
                break
        
        if not placed:
            batches.append(Batch(items=[item], tokens=item.tokens))
    
    return batches


# ============================================================================
# 断点续传管理
# ============================================================================

class PendingFilesManager:
    """待处理文件日志管理器"""
    
    def __init__(self, log_path: Path):
        self._log_path = log_path
    
    def load(self) -> List[Path]:
        """加载待处理文件列表"""
        if not self._log_path.exists():
            return []
        
        try:
            with open(self._log_path, 'r', encoding='utf-8') as f:
                return [Path(line.strip()) for line in f if line.strip()]
        except IOError:
            return []
    
    def save(self, paths: Sequence[Path]) -> None:
        """保存待处理文件列表"""
        try:
            with open(self._log_path, 'w', encoding='utf-8') as f:
                for path in paths:
                    f.write(f"{path}\n")
        except IOError as e:
            logger.error(f"保存待处理文件日志失败: {e}")
    
    def clear(self) -> None:
        """清空日志"""
        if self._log_path.exists():
            try:
                self._log_path.unlink()
                logger.info("待处理文件日志已清空。")
            except OSError as e:
                logger.error(f"清空日志失败: {e}")


# ============================================================================
# 错误处理
# ============================================================================

def is_quota_error(error: Exception) -> bool:
    """判断是否为配额错误"""
    msg = str(error).lower()
    return any(keyword in msg for keyword in ("quota", "exceeded", "429"))


# ============================================================================
# 处理器
# ============================================================================

class FileProcessor:
    """文件处理器"""
    
    def __init__(
        self,
        config: Config,
        limiter: RateLimiter,
        renamer: FileRenamer
    ):
        self._config = config
        self._limiter = limiter
        self._renamer = renamer
    
    async def process_batch(
        self,
        batch: Batch,
        pbar: tqdm
    ) -> BatchResult:
        """处理单个批次"""
        if not batch.items or not MODEL.is_configured:
            pbar.update(len(batch.items))
            return BatchResult(success=False, failed_items=batch.items)
        
        # 构建 prompt
        parts = [PROMPTS['batch']]
        for item in batch.items:
            parts.extend([
                f"\n\n--- START OF FILE: {item.path.name} ---\n",
                item.text,
                f"\n--- END OF FILE: {item.path.name} ---"
            ])
        prompt = "".join(parts)
        
        for attempt in range(self._config.max_retries):
            try:
                await self._limiter.acquire(batch.tokens)
                response = await MODEL.generate_content(prompt, BATCH_SCHEMA)
                results = json.loads(response)
                
                if not isinstance(results, list) or len(results) != len(batch.items):
                    logger.warning(
                        f"批处理结果数不匹配: 预期 {len(batch.items)}, 得到 {len(results)}"
                    )
                    pbar.update(len(batch.items))
                    return BatchResult(success=False, failed_items=batch.items)
                
                # 处理每个文件
                for item, info in zip(batch.items, results):
                    await self._renamer.process(item.path, info)
                
                pbar.update(len(batch.items))
                return BatchResult(success=True)
                
            except json.JSONDecodeError as e:
                logger.error(f"JSON 解析失败: {e}")
                break
            except Exception as e:
                logger.error(f"批处理错误 (尝试 {attempt + 1}): {e}")
                
                if is_quota_error(e):
                    logger.warning("配额已用尽")
                    return BatchResult(
                        success=False,
                        failed_items=batch.items,
                        quota_exceeded=True
                    )
                
                if attempt < self._config.max_retries - 1:
                    await asyncio.sleep(2 ** (attempt + 1))
        
        pbar.update(len(batch.items))
        return BatchResult(success=False, failed_items=batch.items)
    
    async def process_single(
        self,
        item: FileItem,
        pbar: tqdm
    ) -> SingleResult:
        """处理单个文件"""
        if not MODEL.is_configured:
            pbar.update(1)
            return SingleResult(success=False, failed_item=item)
        
        prompt = f"{PROMPTS['single']}\n\n{item.text}"
        
        for attempt in range(self._config.max_retries):
            try:
                await self._limiter.acquire(item.tokens)
                response = await MODEL.generate_content(prompt, SINGLE_OBJECT_SCHEMA)
                info = json.loads(response)
                
                await self._renamer.process(item.path, info)
                
                pbar.update(1)
                return SingleResult(success=True)
                
            except json.JSONDecodeError as e:
                logger.error(f"JSON 解析失败 {item.path.name}: {e}")
                break
            except Exception as e:
                logger.error(f"处理错误 {item.path.name} (尝试 {attempt + 1}): {e}")
                
                if is_quota_error(e):
                    return SingleResult(
                        success=False,
                        failed_item=item,
                        quota_exceeded=True
                    )
                
                if attempt < self._config.max_retries - 1:
                    await asyncio.sleep(2 ** (attempt + 1))
        
        pbar.update(1)
        return SingleResult(success=False, failed_item=item)


# ============================================================================
# 主应用
# ============================================================================

class Application:
    """主应用类"""
    
    def __init__(
        self,
        config: Config,
        mode: ProcessingMode,
        write_metadata: bool,
        target_dir: Path
    ):
        self._config = config
        self._mode = mode
        self._write_metadata = write_metadata
        self._target_dir = target_dir
        
        self._key_manager: Optional[APIKeyManager] = None
        self._pending_manager = PendingFilesManager(config.pending_files_log)
        self._limiter = RateLimiter(config.rpm_limit, config.tpm_limit)
        self._renamer = FileRenamer(write_metadata)
        self._processor = FileProcessor(config, self._limiter, self._renamer)
        self._stats = ProcessingStats()
    
    async def run(self) -> None:
        """运行应用"""
        start_time = time.time()
        
        # 初始化 API 密钥
        api_keys = configure_api_keys()
        self._key_manager = APIKeyManager(api_keys, self._config.tracker_file)
        
        logger.info(f"运行模式: {'批处理' if self._mode == ProcessingMode.BATCH else '单文件'}")
        logger.info(f"元数据写入: {'开启' if self._write_metadata else '关闭'}")
        
        # 确保目标目录存在
        if not self._target_dir.is_dir():
            self._target_dir.mkdir(exist_ok=True)
            logger.info(f"已创建目录: {self._target_dir}")
            return
        
        # 获取待处理文件
        file_paths = self._get_files_to_process()
        if not file_paths:
            logger.info("没有待处理的文件。")
            return
        
        # 准备阶段
        prep_start = time.time()
        file_items = await self._prepare_files(file_paths)
        self._stats.prep_time = time.time() - prep_start
        
        if not file_items:
            logger.warning("未能提取任何文件内容。")
            return
        
        # 处理阶段
        api_start = time.time()
        remaining = await self._process_files(file_items)
        self._stats.api_time = time.time() - api_start
        
        # 保存状态
        if remaining:
            self._pending_manager.save([item.path for item in remaining])
            logger.warning(f"剩余 {len(remaining)} 个文件未处理。")
        else:
            self._pending_manager.clear()
        
        self._print_summary()
    
    def _get_files_to_process(self) -> List[Path]:
        """获取待处理文件列表"""
        pending = self._pending_manager.load()
        
        if pending:
            logger.info(f"从断点恢复: {len(pending)} 个文件")
            return [p for p in pending if p.exists()]
        
        logger.info("扫描目录...")
        paths = []
        for ext in self._config.supported_extensions:
            paths.extend(self._target_dir.glob(f"**/*{ext}"))
        
        return sorted(paths, key=str)
    
    async def _prepare_files(self, paths: List[Path]) -> List[FileItem]:
        """准备文件：提取文本和计算 token"""
        # 先配置一个可用的 API 密钥用于计算 token
        for key in self._key_manager.keys:
            if MODEL.configure(key):
                break
        else:
            logger.error("所有 API 密钥均无效。")
            return []
        
        logger.info(f"提取文本和计算 token (并发={self._config.io_workers})...")
        
        loop = asyncio.get_running_loop()
        items = []
        
        with ThreadPoolExecutor(max_workers=self._config.io_workers) as pool:
            tasks = [
                loop.run_in_executor(pool, extract_and_count, p, self._config)
                for p in paths
            ]
            results = await asyncio.gather(*tasks)
        
        for result in results:
            if result and result.tokens <= self._config.max_tokens_per_request:
                items.append(result)
            elif result:
                logger.warning(f"文件过大，跳过: {result.path.name}")
        
        logger.info(f"准备完成: {len(items)} 个文件")
        return items
    
    async def _process_files(self, items: List[FileItem]) -> List[FileItem]:
        """处理文件"""
        remaining: Deque[FileItem] = deque(items)
        
        for key_index, api_key in enumerate(self._key_manager.keys):
            if not remaining:
                break
            
            logger.info(f"\n--- 第 {key_index + 1}/{self._key_manager.count} 遍 ---")
            logger.info(f"待处理: {len(remaining)} 个文件")
            
            if not MODEL.configure(api_key):
                continue
            
            # 检查配额
            quota = self._key_manager.get_remaining_quota(
                api_key, self._config.daily_request_limit
            )
            if quota <= 0:
                logger.warning("配额已用尽，跳过此密钥。")
                continue
            
            logger.info(f"剩余配额: {quota}")
            
            # 处理
            current_items = list(remaining)
            remaining = deque()
            
            with tqdm(total=len(current_items), desc=f"密钥 #{key_index + 1}", unit="file") as pbar:
                if self._mode == ProcessingMode.BATCH:
                    remaining = await self._process_batch_mode(
                        current_items, quota, pbar, api_key
                    )
                else:
                    remaining = await self._process_single_mode(
                        current_items, quota, pbar, api_key
                    )
            
            self._key_manager.save_tracker()
        
        return list(remaining)
    
    async def _process_batch_mode(
        self,
        items: List[FileItem],
        quota: int,
        pbar: tqdm,
        api_key: str
    ) -> Deque[FileItem]:
        """批处理模式"""
        batches = pack_batches_ffd(
            items,
            self._config.max_tokens_per_request,
            self._config.max_items_per_batch
        )
        
        remaining: Deque[FileItem] = deque()
        
        # 受配额限制的批次
        batches_to_process = batches[:quota]
        leftover_batches = batches[quota:]
        
        for batch in leftover_batches:
            remaining.extend(batch.items)
        
        for batch in batches_to_process:
            self._key_manager.increment_usage(api_key)
            result = await self._processor.process_batch(batch, pbar)
            
            if result.success:
                self._stats.total_processed += len(batch.items)
            else:
                remaining.extend(result.failed_items)
                if result.quota_exceeded:
                    break
        
        return remaining
    
    async def _process_single_mode(
        self,
        items: List[FileItem],
        quota: int,
        pbar: tqdm,
        api_key: str
    ) -> Deque[FileItem]:
        """单文件模式"""
        remaining: Deque[FileItem] = deque()
        processed = 0
        
        for item in items:
            if processed >= quota:
                remaining.append(item)
                continue
            
            self._key_manager.increment_usage(api_key)
            processed += 1
            
            result = await self._processor.process_single(item, pbar)
            
            if result.success:
                self._stats.total_processed += 1
            else:
                if result.failed_item:
                    remaining.append(result.failed_item)
                if result.quota_exceeded:
                    # 将剩余项添加到队列
                    idx = items.index(item)
                    remaining.extend(items[idx + 1:])
                    break
        
        return remaining
    
    def _print_summary(self) -> None:
        """打印运行摘要"""
        print("\n" + "-" * 65)
        print("运行结束！")
        print(f"成功处理: {self._stats.total_processed} 个文件")
        print(f"\n--- 耗时分析 ---")
        print(f"准备阶段: {self._stats.prep_time:.2f} 秒")
        print(f"API处理: {self._stats.api_time:.2f} 秒")
        print(f"总耗时: {self._stats.total_time:.2f} 秒")
        if self._stats.total_processed > 0:
            print(f"平均速率: {self._stats.average_rate:.2f} 文件/秒")


# ============================================================================
# 命令行接口
# ============================================================================

def parse_args() -> argparse.Namespace:
    """解析命令行参数"""
    parser = argparse.ArgumentParser(
        description="使用 Gemini API 批量智能重命名文件并写入元数据。",
        formatter_class=argparse.RawTextHelpFormatter
    )
    
    parser.add_argument(
        "directory",
        nargs='?',
        default="./files_to_rename",
        help="待处理文件目录 (默认: ./files_to_rename)"
    )
    
    parser.add_argument(
        "--mode",
        choices=['batch', 'single'],
        default='batch',
        help="处理模式:\n  batch: 批处理 (默认)\n  single: 单文件处理"
    )
    
    parser.add_argument(
        "--no-metadata",
        action="store_true",
        help="禁用元数据写入"
    )
    
    return parser.parse_args()


async def main() -> None:
    """主入口"""
    args = parse_args()
    
    mode = ProcessingMode.BATCH if args.mode == 'batch' else ProcessingMode.SINGLE
    write_metadata = not args.no_metadata
    target_dir = Path(args.directory)
    
    app = Application(
        config=CONFIG,
        mode=mode,
        write_metadata=write_metadata,
        target_dir=target_dir
    )
    
    await app.run()


if __name__ == "__main__":
    try:
        asyncio.run(main())
    except KeyboardInterrupt:
        print("\n程序被用户中断。")
        sys.exit(0)

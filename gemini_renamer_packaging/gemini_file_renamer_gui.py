# --------------------------------------------------------------------------------
# gemini_file_renamer_gui.py
#
# Gemini File Renaming Tool with GUI / 带图形界面的 Gemini 文件重命名工具
# 
# Features / 功能:
# - Bilingual UI (Chinese/English) / 双语界面支持
# - Modern GUI with drag-and-drop folder selection / 现代化图形界面
# - Dual mode: Batch processing and single file processing / 双模式支持
# - Real-time progress and log display / 实时显示处理进度和日志
# - Multiple API key management / 支持多 API 密钥管理
# - Metadata writing (optional) / 元数据写入功能
# - Resume from breakpoint / 断点续传支持
#
# Dependencies / 依赖安装:
# pip install google-generativeai pymupdf pathvalidate tqdm python-docx EbookLib beautifulsoup4
# --------------------------------------------------------------------------------

from __future__ import annotations

import asyncio
import json
import logging
import os
import sys
import threading
import time
import queue
from abc import ABC, abstractmethod
from collections import deque
from concurrent.futures import ThreadPoolExecutor
from dataclasses import dataclass, field
from datetime import date
from enum import Enum, auto
from pathlib import Path
from typing import (
    Any,
    Callable,
    Deque,
    Dict,
    List,
    Optional,
    Sequence,
    Tuple,
)
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext

# Third-party libraries / 第三方库
try:
    import google.generativeai as genai
    import pymupdf
    from bs4 import BeautifulSoup
    from docx import Document
    from ebooklib import ITEM_DOCUMENT, epub
    from pathvalidate import sanitize_filename
except ImportError as e:
    print(f"Missing dependency / 缺少必要的依赖: {e}")
    print("Please run / 请运行: pip install google-generativeai pymupdf pathvalidate python-docx EbookLib beautifulsoup4")
    sys.exit(1)


# ============================================================================
# Internationalization / 国际化模块
# ============================================================================

class I18n:
    """Internationalization support / 国际化支持"""
    
    LANGUAGES = {
        'zh': {
            'app_title': 'Gemini 智能文件重命名工具',
            'header_title': '📚 Gemini 智能文件重命名',
            'header_subtitle': '使用 AI 自动识别文档内容，智能重命名并写入元数据',
            
            # API Card
            'api_card_title': '🔑 API 密钥设置',
            'api_hint': '输入一个或多个 Google API 密钥（多个密钥用逗号分隔）',
            'show_key': '👁 显示密钥',
            'hide_key': '👁 隐藏密钥',
            
            # Directory Card
            'dir_card_title': '📁 选择目录',
            'browse': '浏览...',
            'select_dir_hint': '请选择包含待处理文件的目录',
            'invalid_dir': '请选择有效的目录',
            'files_found': '找到 {count} 个支持的文件 (PDF, EPUB, AZW3, DOCX)',
            'no_files_found': '未找到支持的文件',
            
            # Options Card
            'options_card_title': '⚙️ 处理选项',
            'processing_mode': '处理模式',
            'batch_mode': '批处理模式（推荐，更高效）',
            'single_mode': '单文件模式（更稳定）',
            'other_options': '其他选项',
            'write_metadata': '写入文件元数据',
            
            # Progress
            'log_title': '📋 处理日志',
            'ready': '就绪',
            'progress': '进度: {current}/{total} ({percent:.1f}%)',
            'preparing': '准备中...',
            'processing': '正在处理...',
            
            # Buttons
            'start': '🚀 开始处理',
            'processing_btn': '处理中...',
            'clear_log': '清除日志',
            
            # Messages
            'error': '错误',
            'api_key_required': '请输入 API 密钥',
            'dir_required': '请选择目标目录',
            'dir_not_exist': '目标目录不存在',
            'api_keys_found': '找到 {count} 个 API 密钥',
            'no_valid_keys': '没有有效的 API 密钥',
            'mode_batch': '批处理',
            'mode_single': '单文件',
            'processing_mode_log': '处理模式: {mode}',
            'metadata_enabled': '开启',
            'metadata_disabled': '关闭',
            'metadata_log': '元数据写入: {status}',
            'resume_from_breakpoint': '从断点恢复: {count} 个文件',
            'files_to_process': '待处理文件: {count} 个',
            'no_files_to_process': '没有需要处理的文件',
            'api_configured': 'API 密钥配置成功',
            'all_keys_invalid': '所有 API 密钥均无效',
            'extracting_text': '正在提取文本...',
            'extracted_files': '成功提取 {count} 个文件',
            'no_content_extracted': '未能提取任何文件内容',
            'using_key': '--- 使用密钥 #{idx}/{total} ---',
            'quota_exhausted': '配额已用尽，跳过',
            'remaining_quota': '剩余配额: {quota}',
            'remaining_files': '剩余 {count} 个文件未处理',
            'completed': '✅ 处理完成！成功处理 {count} 个文件',
            'processed': '✓ 已处理: {name}',
            'batch_mismatch': '批处理结果数不匹配',
            'json_parse_error': 'JSON 解析失败',
            'batch_error': '批处理错误 (尝试 {attempt}): {error}',
            'single_error': '处理错误 {name}: {error}',
            'processing_error': '处理出错: {error}',
            
            # Language
            'language': '🌐 语言',
            'lang_zh': '中文',
            'lang_en': 'English',
        },
        'en': {
            'app_title': 'Gemini Smart File Renamer',
            'header_title': '📚 Gemini Smart File Renamer',
            'header_subtitle': 'Use AI to automatically recognize document content, rename and write metadata',
            
            # API Card
            'api_card_title': '🔑 API Key Settings',
            'api_hint': 'Enter one or more Google API keys (separate multiple keys with commas)',
            'show_key': '👁 Show Key',
            'hide_key': '👁 Hide Key',
            
            # Directory Card
            'dir_card_title': '📁 Select Directory',
            'browse': 'Browse...',
            'select_dir_hint': 'Please select a directory containing files to process',
            'invalid_dir': 'Please select a valid directory',
            'files_found': 'Found {count} supported files (PDF, EPUB, AZW3, DOCX)',
            'no_files_found': 'No supported files found',
            
            # Options Card
            'options_card_title': '⚙️ Processing Options',
            'processing_mode': 'Processing Mode',
            'batch_mode': 'Batch Mode (Recommended, more efficient)',
            'single_mode': 'Single File Mode (More stable)',
            'other_options': 'Other Options',
            'write_metadata': 'Write file metadata',
            
            # Progress
            'log_title': '📋 Processing Log',
            'ready': 'Ready',
            'progress': 'Progress: {current}/{total} ({percent:.1f}%)',
            'preparing': 'Preparing...',
            'processing': 'Processing...',
            
            # Buttons
            'start': '🚀 Start Processing',
            'processing_btn': 'Processing...',
            'clear_log': 'Clear Log',
            
            # Messages
            'error': 'Error',
            'api_key_required': 'Please enter an API key',
            'dir_required': 'Please select a target directory',
            'dir_not_exist': 'Target directory does not exist',
            'api_keys_found': 'Found {count} API keys',
            'no_valid_keys': 'No valid API keys',
            'mode_batch': 'Batch',
            'mode_single': 'Single File',
            'processing_mode_log': 'Processing mode: {mode}',
            'metadata_enabled': 'Enabled',
            'metadata_disabled': 'Disabled',
            'metadata_log': 'Metadata writing: {status}',
            'resume_from_breakpoint': 'Resuming from breakpoint: {count} files',
            'files_to_process': 'Files to process: {count}',
            'no_files_to_process': 'No files to process',
            'api_configured': 'API key configured successfully',
            'all_keys_invalid': 'All API keys are invalid',
            'extracting_text': 'Extracting text...',
            'extracted_files': 'Successfully extracted {count} files',
            'no_content_extracted': 'Failed to extract any file content',
            'using_key': '--- Using key #{idx}/{total} ---',
            'quota_exhausted': 'Quota exhausted, skipping',
            'remaining_quota': 'Remaining quota: {quota}',
            'remaining_files': '{count} files remaining unprocessed',
            'completed': '✅ Completed! Successfully processed {count} files',
            'processed': '✓ Processed: {name}',
            'batch_mismatch': 'Batch result count mismatch',
            'json_parse_error': 'JSON parse error',
            'batch_error': 'Batch error (attempt {attempt}): {error}',
            'single_error': 'Processing error {name}: {error}',
            'processing_error': 'Processing error: {error}',
            
            # Language
            'language': '🌐 Language',
            'lang_zh': '中文',
            'lang_en': 'English',
        }
    }
    
    def __init__(self, lang: str = 'zh'):
        self._lang = lang if lang in self.LANGUAGES else 'zh'
    
    @property
    def lang(self) -> str:
        return self._lang
    
    @lang.setter
    def lang(self, value: str) -> None:
        if value in self.LANGUAGES:
            self._lang = value
    
    def get(self, key: str, **kwargs) -> str:
        """Get translated string with optional formatting"""
        text = self.LANGUAGES.get(self._lang, {}).get(key, key)
        if kwargs:
            try:
                return text.format(**kwargs)
            except (KeyError, ValueError):
                return text
        return text
    
    def toggle(self) -> str:
        """Toggle between languages and return new language code"""
        self._lang = 'en' if self._lang == 'zh' else 'zh'
        return self._lang


# ============================================================================
# Configuration Module / 配置模块
# ============================================================================

@dataclass(frozen=True)
class Config:
    """Immutable global configuration / 不可变的全局配置类"""
    rpm_limit: int = 10
    tpm_limit: int = 250_000
    daily_request_limit: int = 250
    max_tokens_per_request: int = 27_000
    concurrency_limit: int = 10
    max_retries: int = 3
    max_items_per_batch: int = 12
    io_workers: int = field(default_factory=lambda: min(32, max(4, (os.cpu_count() or 8) * 2)))
    chars_per_token: float = 3.5
    supported_extensions: Tuple[str, ...] = ('.pdf', '.epub', '.azw3', '.docx')
    pending_files_log: Path = field(default_factory=lambda: Path("./pending_files.txt"))
    tracker_file: Path = field(default_factory=lambda: Path("./request_tracker.json"))
    
    @property
    def max_chars_per_request(self) -> int:
        return int(self.max_tokens_per_request * self.chars_per_token)


CONFIG = Config()


class ProcessingMode(Enum):
    BATCH = auto()
    SINGLE = auto()


# ============================================================================
# Prompts and Schema / 提示词和 Schema
# ============================================================================

PROMPTS = {
    'batch': """
Analyze the following text, which contains MULTIPLE documents concatenated together.
Each document starts with a "--- START OF FILE: [filename] ---" marker and ends with an "--- END OF FILE: [filename] ---" marker.
For EACH document provided, extract its metadata. Crucially, also extract a list of 3-5 relevant keywords from the document's content.
Return a single JSON array (a list) containing all the extracted JSON objects.
The order of objects in the final list MUST match the order of the documents in the input text.
Do not add any commentary. Only return the JSON array.
""".strip(),
    
    'single': """
Analyze the text from the following document to extract its metadata.
Based on the content, provide a JSON object with the following details.
Crucially, also extract a list of 3-5 relevant keywords from the document's content.
Do not add any commentary. Only return the JSON object.
""".strip()
}

SINGLE_OBJECT_SCHEMA: Dict[str, Any] = {
    "type": "object",
    "properties": {
        "title": {"type": "string"},
        "authors": {"type": "array", "items": {"type": "string"}},
        "keywords": {"type": "array", "items": {"type": "string"}},
        "translators": {"type": "string"},
        "editors": {"type": "string"},
        "publisher_or_journal": {"type": "string"},
        "journal_volume_issue": {"type": "string"},
        "publication_date": {"type": "string"},
        "start_page": {"type": "integer"}
    },
    "required": ["title"]
}

BATCH_SCHEMA: Dict[str, Any] = {
    "type": "array",
    "items": SINGLE_OBJECT_SCHEMA
}


# ============================================================================
# Data Models / 数据模型
# ============================================================================

@dataclass
class FileItem:
    path: Path
    text: str
    tokens: int
    
    def __hash__(self) -> int:
        return hash(self.path)


@dataclass
class BatchResult:
    success: bool
    failed_items: List[FileItem] = field(default_factory=list)
    quota_exceeded: bool = False


@dataclass
class SingleResult:
    success: bool
    failed_item: Optional[FileItem] = None
    quota_exceeded: bool = False


@dataclass
class Batch:
    items: List[FileItem]
    tokens: int


@dataclass
class ProcessingStats:
    total_processed: int = 0
    total_failed: int = 0
    prep_time: float = 0.0
    api_time: float = 0.0
    
    @property
    def total_time(self) -> float:
        return self.prep_time + self.api_time
    
    @property
    def average_rate(self) -> float:
        return self.total_processed / self.api_time if self.api_time > 0 else 0.0


# ============================================================================
# GUI Log Handler / GUI 日志处理器
# ============================================================================

class QueueHandler(logging.Handler):
    """Handler that sends log messages to a queue / 将日志消息发送到队列的处理器"""
    
    def __init__(self, log_queue: queue.Queue):
        super().__init__()
        self.log_queue = log_queue
    
    def emit(self, record: logging.LogRecord) -> None:
        self.log_queue.put(self.format(record))


# ============================================================================
# API Key Management / API 密钥管理
# ============================================================================

class APIKeyManager:
    def __init__(self, keys: List[str], tracker_file: Path):
        self._keys = keys
        self._tracker_file = tracker_file
        self._tracker = self._load_tracker()
    
    def _load_tracker(self) -> Dict[str, Any]:
        today_str = date.today().isoformat()
        default = {"date": today_str, "usage": {}}
        
        if not self._tracker_file.exists():
            return default
            
        try:
            with open(self._tracker_file, 'r', encoding='utf-8') as f:
                tracker = json.load(f)
            if tracker.get("date") != today_str:
                return default
            tracker.setdefault("usage", {})
            return tracker
        except (json.JSONDecodeError, IOError):
            return default
    
    def save_tracker(self) -> None:
        try:
            with open(self._tracker_file, 'w', encoding='utf-8') as f:
                json.dump(self._tracker, f, indent=4, ensure_ascii=False)
        except IOError:
            pass
    
    def get_usage(self, key: str) -> int:
        return self._tracker["usage"].get(key, 0)
    
    def increment_usage(self, key: str) -> None:
        self._tracker["usage"][key] = self.get_usage(key) + 1
    
    @property
    def keys(self) -> List[str]:
        return self._keys
    
    @property
    def count(self) -> int:
        return len(self._keys)
    
    def get_remaining_quota(self, key: str, daily_limit: int) -> int:
        return daily_limit - self.get_usage(key)


# ============================================================================
# Gemini Model / Gemini 模型
# ============================================================================

class GeminiModel:
    MODEL_NAME = 'models/gemini-2.5-flash'
    
    def __init__(self):
        self._model: Optional[genai.GenerativeModel] = None
        self._api_key: Optional[str] = None
    
    @property
    def api_key(self) -> Optional[str]:
        return self._api_key
    
    @property
    def is_configured(self) -> bool:
        return self._model is not None
    
    def configure(self, api_key: str) -> bool:
        try:
            genai.configure(api_key=api_key)
            self._model = genai.GenerativeModel(self.MODEL_NAME)
            self._api_key = api_key
            return True
        except Exception:
            return False
    
    def count_tokens(self, text: str) -> int:
        if not self._model:
            raise RuntimeError("Model not configured / 模型尚未配置")
        result = self._model.count_tokens(text)
        return result.total_tokens
    
    async def generate_content(self, prompt: str, schema: Dict[str, Any]) -> str:
        if not self._model:
            raise RuntimeError("Model not configured / 模型尚未配置")
        config = {"response_mime_type": "application/json", "response_schema": schema}
        response = await self._model.generate_content_async(prompt, generation_config=config)
        return response.text


MODEL = GeminiModel()


# ============================================================================
# Rate Limiter / 速率限制器
# ============================================================================

class RateLimiter:
    def __init__(self, rpm: int, tpm: int):
        self._rpm = rpm
        self._tpm = tpm
        self._request_timestamps: Deque[float] = deque()
        self._token_records: Deque[Tuple[float, int]] = deque()
        self._token_total = 0
        self._lock = asyncio.Lock()
    
    def _cleanup_old_records(self, now: float) -> None:
        cutoff = now - 60
        while self._request_timestamps and self._request_timestamps[0] < cutoff:
            self._request_timestamps.popleft()
        while self._token_records and self._token_records[0][0] < cutoff:
            _, tokens = self._token_records.popleft()
            self._token_total -= tokens
    
    def _calculate_wait_time(self, now: float, tokens_needed: int) -> float:
        rpm_wait = 0.0
        tpm_wait = 0.0
        
        if len(self._request_timestamps) >= self._rpm and self._request_timestamps:
            rpm_wait = (self._request_timestamps[0] + 60) - now
        
        if (self._token_total + tokens_needed) > self._tpm and self._token_records:
            tokens_to_free = (self._token_total + tokens_needed) - self._tpm
            freed = 0
            wait_until = 0.0
            for ts, tk in self._token_records:
                freed += tk
                if freed >= tokens_to_free:
                    wait_until = ts
                    break
            if wait_until > 0:
                tpm_wait = (wait_until + 60) - now
        
        return max(0.1, rpm_wait, tpm_wait)
    
    async def acquire(self, tokens_needed: int) -> None:
        async with self._lock:
            while True:
                now = time.time()
                self._cleanup_old_records(now)
                
                if (len(self._request_timestamps) < self._rpm and 
                    (self._token_total + tokens_needed) <= self._tpm):
                    self._request_timestamps.append(now)
                    self._token_records.append((now, tokens_needed))
                    self._token_total += tokens_needed
                    return
                
                wait_time = self._calculate_wait_time(now, tokens_needed)
                await asyncio.sleep(wait_time)


# ============================================================================
# Text Extractors / 文本提取器
# ============================================================================

class TextExtractor(ABC):
    @abstractmethod
    def extract(self, path: Path) -> str:
        pass


class PDFExtractor(TextExtractor):
    def __init__(self, pages_start: int = 4, pages_end: int = 3):
        self._pages_start = pages_start
        self._pages_end = pages_end
    
    def extract(self, path: Path) -> str:
        try:
            with pymupdf.open(path) as doc:
                total = doc.page_count
                pages = set(range(min(self._pages_start, total)))
                if total > self._pages_start + self._pages_end:
                    pages.update(range(total - self._pages_end, total))
                texts = [doc[i].get_text(sort=True) for i in sorted(pages)]
                return "\n".join(texts)
        except Exception:
            return ""


class EPUBExtractor(TextExtractor):
    def __init__(self, chapters_start: int = 5, chapters_end: int = 4):
        self._chapters_start = chapters_start
        self._chapters_end = chapters_end
    
    def extract(self, path: Path) -> str:
        try:
            book = epub.read_epub(path)
            items = list(book.get_items_of_type(ITEM_DOCUMENT))
            to_process = items[:self._chapters_start]
            if len(items) > self._chapters_start + self._chapters_end:
                to_process.extend(items[-self._chapters_end:])
            texts = []
            for item in to_process:
                soup = BeautifulSoup(item.get_body_content(), 'html.parser')
                texts.append(soup.get_text("\n", strip=True))
            return "\n\n".join(texts)
        except Exception:
            return ""


class DOCXExtractor(TextExtractor):
    def __init__(self, paras_start: int = 20, paras_end: int = 15):
        self._paras_start = paras_start
        self._paras_end = paras_end
    
    def extract(self, path: Path) -> str:
        try:
            doc = Document(path)
            paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            if len(paras) > self._paras_start + self._paras_end:
                result = paras[:self._paras_start] + paras[-self._paras_end:]
            else:
                result = paras
            return "\n".join(result)
        except Exception:
            return ""


class TextExtractorFactory:
    _extractors: Dict[str, TextExtractor] = {
        '.pdf': PDFExtractor(),
        '.epub': EPUBExtractor(),
        '.azw3': EPUBExtractor(),
        '.docx': DOCXExtractor(),
    }
    
    @classmethod
    def get_extractor(cls, extension: str) -> Optional[TextExtractor]:
        return cls._extractors.get(extension.lower())


def smart_truncate(text: str, max_chars: int) -> str:
    if len(text) <= max_chars:
        return text
    head_chars = int(max_chars * 0.6)
    tail_chars = int(max_chars * 0.4)
    return f"{text[:head_chars]}\n\n--- Content truncated / 内容已截断 ---\n\n{text[-tail_chars:]}"


def extract_text(path: Path, max_chars: int) -> Optional[str]:
    extractor = TextExtractorFactory.get_extractor(path.suffix)
    if not extractor:
        return None
    text = extractor.extract(path)
    if not text:
        return None
    return smart_truncate(text, max_chars)


def extract_and_count(path: Path, config: Config) -> Optional[FileItem]:
    text = extract_text(path, config.max_chars_per_request)
    if not text:
        return None
    if not MODEL.is_configured:
        return None
    try:
        tokens = MODEL.count_tokens(text)
        return FileItem(path=path, text=text, tokens=tokens)
    except Exception:
        return None


# ============================================================================
# Metadata Processing / 元数据处理
# ============================================================================

JOURNAL_KEYWORDS = frozenset([
    "journal", "review", "proceedings", "transactions", "quarterly",
    "annals", "bulletin", "magazine", "advances", "letters", "studies",
    "science", "research", "technology", "medicine", "report", "archives",
    "学报", "法学", "研究", "评论", "科学", "技术", "杂志", "动态",
    "报告", "医学", "经济", "哲学", "历史", "通讯", "汇刊", "纪要"
])

UNKNOWN_AUTHOR_MARKERS = frozenset(["作者不详"])
ROLE_INVALID_TOKENS = frozenset(["null", "none", "n/a", "unknown", "不详", "未知"])
ROLE_INVALID_SUBSTRINGS = frozenset(["无法提取", "不明确", "系统返回null", "系统返回 null"])


class TextNormalizer:
    @staticmethod
    def normalize(value: Any) -> str:
        if value is None:
            return ""
        text = str(value).strip()
        return "" if not text or text.lower() == "null" else text
    
    @staticmethod
    def normalize_list(values: Optional[List[Any]]) -> List[str]:
        if not values:
            return []
        return [v for v in map(TextNormalizer.normalize, values) if v]
    
    @staticmethod
    def normalize_authors(values: Optional[List[Any]]) -> List[str]:
        authors = TextNormalizer.normalize_list(values)
        return [a for a in authors if a not in UNKNOWN_AUTHOR_MARKERS]
    
    @staticmethod
    def normalize_role(value: Any) -> str:
        normalized = TextNormalizer.normalize(value)
        if not normalized:
            return ""
        lower = normalized.lower()
        if lower in ROLE_INVALID_TOKENS:
            return ""
        compact = normalized.replace(" ", "")
        for marker in ROLE_INVALID_SUBSTRINGS:
            if marker in normalized or marker in compact:
                return ""
        return normalized


class MetadataBuilder:
    def __init__(self, info: Dict[str, Any]):
        self._info = info
        self._normalizer = TextNormalizer
    
    @property
    def title(self) -> str:
        return self._normalizer.normalize(self._info.get("title"))
    
    @property
    def authors(self) -> List[str]:
        return self._normalizer.normalize_authors(self._info.get("authors"))
    
    @property
    def authors_str(self) -> str:
        return "、".join(self.authors)
    
    @property
    def keywords(self) -> List[str]:
        return self._normalizer.normalize_list(self._info.get("keywords"))
    
    @property
    def keywords_str(self) -> str:
        return ", ".join(self.keywords)
    
    @property
    def translators(self) -> str:
        return self._normalizer.normalize_role(self._info.get("translators"))
    
    @property
    def editors(self) -> str:
        return self._normalizer.normalize_role(self._info.get("editors"))
    
    @property
    def publisher(self) -> str:
        return self._normalizer.normalize(self._info.get("publisher_or_journal"))
    
    def build_details_string(self) -> str:
        details = []
        mappings = [
            ("Publisher/Journal", self.publisher),
            ("Volume/Issue", self._normalizer.normalize(self._info.get("journal_volume_issue"))),
            ("Date", self._normalizer.normalize(self._info.get("publication_date"))),
            ("Editor", self.editors),
            ("Translator", self.translators),
            ("Page", self._normalizer.normalize(self._info.get("start_page"))),
        ]
        for label, value in mappings:
            if value:
                details.append(f"{label}: {value}")
        return " | ".join(details)
    
    def build_filename(self) -> Optional[str]:
        """Build filename"""
        if not self.title:
            return None
        
        main_part = f"{self.title} - {self.authors_str}" if self.authors_str else self.title
        
        extras = []
        if self.translators:
            extras.append(f"{self.translators} 译")
        
        if self.editors and not self.authors:
            pub_lower = self.publisher.lower()
            if not any(k in pub_lower for k in JOURNAL_KEYWORDS):
                extras.append(f"{self.editors} 编")
        
        return f"{main_part} ({', '.join(extras)})" if extras else main_part


# ============================================================================
# Metadata Writers / 元数据写入器
# ============================================================================

class MetadataWriter(ABC):
    @abstractmethod
    def write(self, path: Path, builder: MetadataBuilder) -> None:
        pass


class PDFMetadataWriter(MetadataWriter):
    def write(self, path: Path, builder: MetadataBuilder) -> None:
        try:
            with pymupdf.open(path) as doc:
                metadata = doc.metadata
                metadata['title'] = builder.title
                metadata['author'] = builder.authors_str
                metadata['subject'] = builder.build_details_string()
                metadata['keywords'] = builder.keywords_str
                doc.set_metadata(metadata)
                doc.save(doc.name, incremental=True, encryption=pymupdf.PDF_ENCRYPT_KEEP)
        except Exception:
            pass


class DOCXMetadataWriter(MetadataWriter):
    def write(self, path: Path, builder: MetadataBuilder) -> None:
        try:
            doc = Document(path)
            cp = doc.core_properties
            cp.title = builder.title
            cp.author = builder.authors_str
            cp.subject = builder.build_details_string()
            cp.keywords = builder.keywords_str
            cp.comments = "Metadata updated by Gemini File Renamer"
            doc.save(path)
        except Exception:
            pass


class EPUBMetadataWriter(MetadataWriter):
    def _clear_creators(self, book: epub.EpubBook) -> None:
        namespace = "http://purl.org/dc/elements/1.1/"
        meta = book.metadata.get(namespace)
        if isinstance(meta, dict) and "creator" in meta:
            meta["creator"] = []
        elif isinstance(meta, list):
            book.metadata[namespace] = [
                item for item in meta 
                if not (isinstance(item, tuple) and item and item[0] == "creator")
            ]
    
    def write(self, path: Path, builder: MetadataBuilder) -> None:
        try:
            book = epub.read_epub(path)
            book.set_title(builder.title)
            self._clear_creators(book)
            for author in builder.authors:
                book.add_author(author)
            description_parts = []
            details = builder.build_details_string()
            if details:
                description_parts.append(details)
            if builder.keywords:
                description_parts.append(f"Keywords: {builder.keywords_str}")
            if description_parts:
                book.add_metadata('DC', 'description', "\n".join(description_parts))
            epub.write_epub(path, book)
        except Exception:
            pass


class MetadataWriterFactory:
    _writers: Dict[str, MetadataWriter] = {
        '.pdf': PDFMetadataWriter(),
        '.docx': DOCXMetadataWriter(),
        '.epub': EPUBMetadataWriter(),
        '.azw3': EPUBMetadataWriter(),
    }
    
    @classmethod
    def get_writer(cls, extension: str) -> Optional[MetadataWriter]:
        return cls._writers.get(extension.lower())


# ============================================================================
# File Renamer / 文件重命名器
# ============================================================================

class FileRenamer:
    def __init__(self, write_metadata: bool = True):
        self._write_metadata = write_metadata
        self._executor = ThreadPoolExecutor(max_workers=4)
    
    async def process(self, path: Path, info: Dict[str, Any]) -> None:
        builder = MetadataBuilder(info)
        new_name = builder.build_filename()
        
        if not new_name:
            return
        
        safe_name = sanitize_filename(new_name).strip()
        if not safe_name or safe_name in {".", ".."}:
            return
        
        new_path = path.with_name(f"{safe_name}{path.suffix}")
        
        counter = 1
        while new_path.exists() and new_path != path:
            new_path = path.with_name(f"{safe_name}_{counter}{path.suffix}")
            counter += 1
        
        if new_path != path:
            try:
                path.rename(new_path)
            except OSError:
                return
        
        if self._write_metadata:
            await self._write_metadata_async(new_path, builder)
    
    async def _write_metadata_async(self, path: Path, builder: MetadataBuilder) -> None:
        writer = MetadataWriterFactory.get_writer(path.suffix)
        if not writer:
            return
        loop = asyncio.get_running_loop()
        await loop.run_in_executor(self._executor, writer.write, path, builder)


# ============================================================================
# Batch Packing / 批处理打包
# ============================================================================

def pack_batches_ffd(
    items: List[FileItem],
    max_tokens: int,
    max_items: Optional[int] = None
) -> List[Batch]:
    sorted_items = sorted(items, key=lambda x: x.tokens, reverse=True)
    batches: List[Batch] = []
    
    for item in sorted_items:
        placed = False
        for batch in batches:
            can_fit_tokens = batch.tokens + item.tokens <= max_tokens
            can_fit_items = max_items is None or len(batch.items) < max_items
            if can_fit_tokens and can_fit_items:
                batch.items.append(item)
                batch.tokens += item.tokens
                placed = True
                break
        if not placed:
            batches.append(Batch(items=[item], tokens=item.tokens))
    
    return batches


# ============================================================================
# Breakpoint Resume / 断点续传
# ============================================================================

class PendingFilesManager:
    def __init__(self, log_path: Path):
        self._log_path = log_path
    
    def load(self) -> List[Path]:
        if not self._log_path.exists():
            return []
        try:
            with open(self._log_path, 'r', encoding='utf-8') as f:
                return [Path(line.strip()) for line in f if line.strip()]
        except IOError:
            return []
    
    def save(self, paths: Sequence[Path]) -> None:
        try:
            with open(self._log_path, 'w', encoding='utf-8') as f:
                for path in paths:
                    f.write(f"{path}\n")
        except IOError:
            pass
    
    def clear(self) -> None:
        if self._log_path.exists():
            try:
                self._log_path.unlink()
            except OSError:
                pass


def is_quota_error(error: Exception) -> bool:
    msg = str(error).lower()
    return any(keyword in msg for keyword in ("quota", "exceeded", "429"))


# ============================================================================
# File Processor / 文件处理器
# ============================================================================

class FileProcessor:
    def __init__(
        self,
        config: Config,
        limiter: RateLimiter,
        renamer: FileRenamer,
        logger: logging.Logger,
        i18n: I18n
    ):
        self._config = config
        self._limiter = limiter
        self._renamer = renamer
        self._logger = logger
        self._i18n = i18n
    
    async def process_batch(
        self,
        batch: Batch,
        progress_callback: Optional[Callable[[int], None]] = None
    ) -> BatchResult:
        if not batch.items or not MODEL.is_configured:
            if progress_callback:
                progress_callback(len(batch.items))
            return BatchResult(success=False, failed_items=batch.items)
        
        parts = [PROMPTS['batch']]
        for item in batch.items:
            parts.extend([
                f"\n\n--- START OF FILE: {item.path.name} ---\n",
                item.text,
                f"\n--- END OF FILE: {item.path.name} ---"
            ])
        prompt = "".join(parts)
        
        for attempt in range(self._config.max_retries):
            try:
                await self._limiter.acquire(batch.tokens)
                response = await MODEL.generate_content(prompt, BATCH_SCHEMA)
                results = json.loads(response)
                
                if not isinstance(results, list) or len(results) != len(batch.items):
                    self._logger.warning(self._i18n.get('batch_mismatch'))
                    if progress_callback:
                        progress_callback(len(batch.items))
                    return BatchResult(success=False, failed_items=batch.items)
                
                for item, info in zip(batch.items, results):
                    await self._renamer.process(item.path, info)
                    self._logger.info(self._i18n.get('processed', name=item.path.name))
                
                if progress_callback:
                    progress_callback(len(batch.items))
                return BatchResult(success=True)
                
            except json.JSONDecodeError:
                self._logger.error(self._i18n.get('json_parse_error'))
                break
            except Exception as e:
                self._logger.error(self._i18n.get('batch_error', attempt=attempt + 1, error=e))
                if is_quota_error(e):
                    return BatchResult(success=False, failed_items=batch.items, quota_exceeded=True)
                if attempt < self._config.max_retries - 1:
                    await asyncio.sleep(2 ** (attempt + 1))
        
        if progress_callback:
            progress_callback(len(batch.items))
        return BatchResult(success=False, failed_items=batch.items)
    
    async def process_single(
        self,
        item: FileItem,
        progress_callback: Optional[Callable[[int], None]] = None
    ) -> SingleResult:
        if not MODEL.is_configured:
            if progress_callback:
                progress_callback(1)
            return SingleResult(success=False, failed_item=item)
        
        prompt = f"{PROMPTS['single']}\n\n{item.text}"
        
        for attempt in range(self._config.max_retries):
            try:
                await self._limiter.acquire(item.tokens)
                response = await MODEL.generate_content(prompt, SINGLE_OBJECT_SCHEMA)
                info = json.loads(response)
                
                await self._renamer.process(item.path, info)
                self._logger.info(self._i18n.get('processed', name=item.path.name))
                
                if progress_callback:
                    progress_callback(1)
                return SingleResult(success=True)
                
            except json.JSONDecodeError:
                self._logger.error(self._i18n.get('json_parse_error') + f": {item.path.name}")
                break
            except Exception as e:
                self._logger.error(self._i18n.get('single_error', name=item.path.name, error=e))
                if is_quota_error(e):
                    return SingleResult(success=False, failed_item=item, quota_exceeded=True)
                if attempt < self._config.max_retries - 1:
                    await asyncio.sleep(2 ** (attempt + 1))
        
        if progress_callback:
            progress_callback(1)
        return SingleResult(success=False, failed_item=item)


# ============================================================================
# Modern GUI / 现代化 GUI 界面
# ============================================================================

class ModernStyle:
    """Modern style definitions / 现代化样式定义"""
    
    # Color scheme - Dark theme / 颜色方案 - 深色主题
    BG_PRIMARY = "#1a1a2e"
    BG_SECONDARY = "#16213e"
    BG_TERTIARY = "#0f3460"
    ACCENT = "#e94560"
    ACCENT_HOVER = "#ff6b6b"
    TEXT_PRIMARY = "#eaeaea"
    TEXT_SECONDARY = "#a0a0a0"
    SUCCESS = "#4ecca3"
    WARNING = "#ffc107"
    ERROR = "#ff6b6b"
    BORDER = "#2d3a4f"
    
    # Fonts / 字体
    FONT_FAMILY = "Segoe UI"
    FONT_FAMILY_FALLBACK = ("Segoe UI", "PingFang SC", "Microsoft YaHei", "Helvetica", "Arial")
    FONT_SIZE_LARGE = 13
    FONT_SIZE_NORMAL = 11
    FONT_SIZE_SMALL = 10
    
    @classmethod
    def configure_styles(cls, root: tk.Tk) -> None:
        """Configure ttk styles / 配置 ttk 样式"""
        style = ttk.Style(root)
        
        try:
            style.theme_use('clam')
        except tk.TclError:
            pass
        
        style.configure("Modern.TFrame", background=cls.BG_PRIMARY)
        style.configure("Card.TFrame", background=cls.BG_SECONDARY, relief="flat")
        style.configure("Modern.TLabel", background=cls.BG_PRIMARY, foreground=cls.TEXT_PRIMARY,
                       font=(cls.FONT_FAMILY, cls.FONT_SIZE_NORMAL))
        style.configure("Title.TLabel", background=cls.BG_PRIMARY, foreground=cls.TEXT_PRIMARY,
                       font=(cls.FONT_FAMILY, 20, "bold"))
        style.configure("Subtitle.TLabel", background=cls.BG_PRIMARY, foreground=cls.TEXT_SECONDARY,
                       font=(cls.FONT_FAMILY, cls.FONT_SIZE_SMALL))
        style.configure("CardTitle.TLabel", background=cls.BG_SECONDARY, foreground=cls.TEXT_PRIMARY,
                       font=(cls.FONT_FAMILY, cls.FONT_SIZE_LARGE, "bold"))
        style.configure("CardText.TLabel", background=cls.BG_SECONDARY, foreground=cls.TEXT_SECONDARY,
                       font=(cls.FONT_FAMILY, cls.FONT_SIZE_NORMAL))
        style.configure("Accent.TButton", background=cls.ACCENT, foreground="white",
                       font=(cls.FONT_FAMILY, cls.FONT_SIZE_NORMAL, "bold"), padding=(20, 12), borderwidth=0)
        style.map("Accent.TButton", background=[("active", cls.ACCENT_HOVER), ("disabled", cls.BG_TERTIARY)],
                 foreground=[("disabled", cls.TEXT_SECONDARY)])
        style.configure("Secondary.TButton", background=cls.BG_TERTIARY, foreground=cls.TEXT_PRIMARY,
                       font=(cls.FONT_FAMILY, cls.FONT_SIZE_NORMAL), padding=(16, 10), borderwidth=0)
        style.map("Secondary.TButton", background=[("active", cls.BORDER)])
        style.configure("Modern.Horizontal.TProgressbar", background=cls.ACCENT, troughcolor=cls.BG_TERTIARY,
                       borderwidth=0, lightcolor=cls.ACCENT, darkcolor=cls.ACCENT)
        style.configure("Modern.TCheckbutton", background=cls.BG_SECONDARY, foreground=cls.TEXT_PRIMARY,
                       font=(cls.FONT_FAMILY, cls.FONT_SIZE_NORMAL), indicatorbackground=cls.BG_TERTIARY,
                       indicatorforeground=cls.ACCENT)
        style.map("Modern.TCheckbutton", background=[("active", cls.BG_SECONDARY)],
                 indicatorbackground=[("selected", cls.ACCENT)])
        style.configure("Modern.TRadiobutton", background=cls.BG_SECONDARY, foreground=cls.TEXT_PRIMARY,
                       font=(cls.FONT_FAMILY, cls.FONT_SIZE_NORMAL), indicatorbackground=cls.BG_TERTIARY)
        style.map("Modern.TRadiobutton", background=[("active", cls.BG_SECONDARY)],
                 indicatorbackground=[("selected", cls.ACCENT)])


class GeminiRenamerGUI:
    """Gemini File Renamer GUI / Gemini 文件重命名工具 GUI"""
    
    def __init__(self):
        self.root = tk.Tk()
        self.i18n = I18n('zh')
        
        self.root.title(self.i18n.get('app_title'))
        self.root.geometry("900x750")
        self.root.minsize(800, 650)
        self.root.configure(bg=ModernStyle.BG_PRIMARY)
        
        ModernStyle.configure_styles(self.root)
        
        # State variables / 状态变量
        self.target_dir = tk.StringVar(value="")
        self.api_keys = tk.StringVar(value=os.getenv("GOOGLE_API_KEY", ""))
        self.mode = tk.StringVar(value="batch")
        self.write_metadata = tk.BooleanVar(value=True)
        self.is_processing = False
        self.show_key = tk.BooleanVar(value=False)
        
        # Log queue / 日志队列
        self.log_queue: queue.Queue = queue.Queue()
        
        # Configure logging / 配置日志
        self.logger = logging.getLogger("GeminiRenamer")
        self.logger.setLevel(logging.INFO)
        handler = QueueHandler(self.log_queue)
        handler.setFormatter(logging.Formatter('%(asctime)s - %(message)s', datefmt='%H:%M:%S'))
        self.logger.addHandler(handler)
        
        # Store widget references for language updates / 存储组件引用用于语言更新
        self.widgets: Dict[str, Any] = {}
        
        self._build_ui()
        self._update_log()
    
    def _build_ui(self) -> None:
        """Build user interface / 构建用户界面"""
        # Main scrollable container / 主滚动容器
        self.canvas = tk.Canvas(self.root, bg=ModernStyle.BG_PRIMARY, highlightthickness=0)
        scrollbar = ttk.Scrollbar(self.root, orient="vertical", command=self.canvas.yview)
        self.scrollable_frame = ttk.Frame(self.canvas, style="Modern.TFrame")
        
        self.scrollable_frame.bind(
            "<Configure>",
            lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all"))
        )
        
        self.canvas_window = self.canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw")
        self.canvas.configure(yscrollcommand=scrollbar.set)
        
        # Bind mouse wheel / 绑定鼠标滚轮
        self.canvas.bind_all("<MouseWheel>", self._on_mousewheel)
        self.canvas.bind_all("<Button-4>", self._on_mousewheel)
        self.canvas.bind_all("<Button-5>", self._on_mousewheel)
        
        # Bind canvas resize / 绑定画布大小调整
        self.canvas.bind("<Configure>", self._on_canvas_configure)
        
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # Main content frame / 主内容框架
        main_frame = ttk.Frame(self.scrollable_frame, style="Modern.TFrame")
        main_frame.pack(fill=tk.BOTH, expand=True, padx=25, pady=15)
        
        self._build_header(main_frame)
        self._build_api_card(main_frame)
        self._build_directory_card(main_frame)
        self._build_options_card(main_frame)
        self._build_progress_section(main_frame)
        self._build_action_buttons(main_frame)
    
    def _on_canvas_configure(self, event: tk.Event) -> None:
        """Handle canvas resize / 处理画布大小调整"""
        self.canvas.itemconfig(self.canvas_window, width=event.width)
    
    def _on_mousewheel(self, event: tk.Event) -> None:
        """Handle mouse wheel scroll / 处理鼠标滚轮滚动"""
        if event.num == 5 or event.delta < 0:
            self.canvas.yview_scroll(1, "units")
        elif event.num == 4 or event.delta > 0:
            self.canvas.yview_scroll(-1, "units")
    
    def _build_header(self, parent: ttk.Frame) -> None:
        """Build header section / 构建标题区域"""
        header_frame = ttk.Frame(parent, style="Modern.TFrame")
        header_frame.pack(fill=tk.X, pady=(0, 15))
        
        # Top row with title and language switch / 顶部行包含标题和语言切换
        top_row = ttk.Frame(header_frame, style="Modern.TFrame")
        top_row.pack(fill=tk.X)
        
        self.widgets['title'] = ttk.Label(top_row, text=self.i18n.get('header_title'), style="Title.TLabel")
        self.widgets['title'].pack(side=tk.LEFT, anchor="w")
        
        # Language switch button / 语言切换按钮
        lang_frame = tk.Frame(top_row, bg=ModernStyle.BG_PRIMARY)
        lang_frame.pack(side=tk.RIGHT)
        
        self.widgets['lang_btn'] = tk.Button(
            lang_frame,
            text=self.i18n.get('language') + ": " + ("中文" if self.i18n.lang == 'zh' else "English"),
            command=self._toggle_language,
            bg=ModernStyle.BG_TERTIARY, fg=ModernStyle.TEXT_PRIMARY,
            activebackground=ModernStyle.BORDER, activeforeground=ModernStyle.TEXT_PRIMARY,
            relief="flat", cursor="hand2",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_SMALL), padx=12, pady=5
        )
        self.widgets['lang_btn'].pack()
        
        self.widgets['subtitle'] = ttk.Label(header_frame, text=self.i18n.get('header_subtitle'), style="Subtitle.TLabel")
        self.widgets['subtitle'].pack(anchor="w", pady=(5, 0))
    
    def _build_api_card(self, parent: ttk.Frame) -> None:
        """Build API key card / 构建 API 密钥卡片"""
        card, content = self._create_card(parent, 'api_card_title')
        
        self.widgets['api_hint'] = ttk.Label(content, text=self.i18n.get('api_hint'), style="CardText.TLabel")
        self.widgets['api_hint'].pack(anchor="w", pady=(0, 8))
        
        entry_frame = tk.Frame(content, bg=ModernStyle.BG_TERTIARY, highlightthickness=1,
                               highlightbackground=ModernStyle.BORDER)
        entry_frame.pack(fill=tk.X)
        
        self.api_entry = tk.Entry(
            entry_frame, textvariable=self.api_keys,
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_NORMAL),
            bg=ModernStyle.BG_TERTIARY, fg=ModernStyle.TEXT_PRIMARY,
            insertbackground=ModernStyle.TEXT_PRIMARY, relief="flat", show="•"
        )
        self.api_entry.pack(fill=tk.X, padx=12, pady=10)
        
        self.widgets['show_key_btn'] = tk.Button(
            content, text=self.i18n.get('show_key'), command=self._toggle_key_visibility,
            bg=ModernStyle.BG_SECONDARY, fg=ModernStyle.TEXT_SECONDARY,
            activebackground=ModernStyle.BG_TERTIARY, activeforeground=ModernStyle.TEXT_PRIMARY,
            relief="flat", cursor="hand2", font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_SMALL)
        )
        self.widgets['show_key_btn'].pack(anchor="w", pady=(8, 0))
    
    def _build_directory_card(self, parent: ttk.Frame) -> None:
        """Build directory selection card / 构建目录选择卡片"""
        card, content = self._create_card(parent, 'dir_card_title')
        
        dir_frame = ttk.Frame(content, style="Card.TFrame")
        dir_frame.pack(fill=tk.X)
        
        path_frame = tk.Frame(dir_frame, bg=ModernStyle.BG_TERTIARY, highlightthickness=1,
                              highlightbackground=ModernStyle.BORDER)
        path_frame.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        self.path_entry = tk.Entry(
            path_frame, textvariable=self.target_dir,
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_NORMAL),
            bg=ModernStyle.BG_TERTIARY, fg=ModernStyle.TEXT_PRIMARY,
            insertbackground=ModernStyle.TEXT_PRIMARY, relief="flat"
        )
        self.path_entry.pack(fill=tk.X, padx=12, pady=10)
        
        self.widgets['browse_btn'] = tk.Button(
            dir_frame, text=self.i18n.get('browse'), command=self._browse_directory,
            bg=ModernStyle.BG_TERTIARY, fg=ModernStyle.TEXT_PRIMARY,
            activebackground=ModernStyle.BORDER, activeforeground=ModernStyle.TEXT_PRIMARY,
            relief="flat", cursor="hand2",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_NORMAL), padx=20, pady=8
        )
        self.widgets['browse_btn'].pack(side=tk.RIGHT, padx=(10, 0))
        
        self.widgets['file_count'] = ttk.Label(content, text=self.i18n.get('select_dir_hint'), style="CardText.TLabel")
        self.widgets['file_count'].pack(anchor="w", pady=(10, 0))
    
    def _build_options_card(self, parent: ttk.Frame) -> None:
        """Build options card / 构建选项卡片"""
        card, content = self._create_card(parent, 'options_card_title')
        
        options_frame = ttk.Frame(content, style="Card.TFrame")
        options_frame.pack(fill=tk.X)
        
        # Left: Processing mode / 左侧：处理模式
        left_frame = ttk.Frame(options_frame, style="Card.TFrame")
        left_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        self.widgets['mode_label'] = ttk.Label(left_frame, text=self.i18n.get('processing_mode'), style="CardText.TLabel")
        self.widgets['mode_label'].pack(anchor="w")
        
        self.widgets['batch_radio'] = ttk.Radiobutton(
            left_frame, text=self.i18n.get('batch_mode'),
            variable=self.mode, value="batch", style="Modern.TRadiobutton"
        )
        self.widgets['batch_radio'].pack(anchor="w", pady=(5, 2))
        
        self.widgets['single_radio'] = ttk.Radiobutton(
            left_frame, text=self.i18n.get('single_mode'),
            variable=self.mode, value="single", style="Modern.TRadiobutton"
        )
        self.widgets['single_radio'].pack(anchor="w")
        
        # Right: Other options / 右侧：其他选项
        right_frame = ttk.Frame(options_frame, style="Card.TFrame")
        right_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=(30, 0))
        
        self.widgets['other_label'] = ttk.Label(right_frame, text=self.i18n.get('other_options'), style="CardText.TLabel")
        self.widgets['other_label'].pack(anchor="w")
        
        self.widgets['metadata_check'] = ttk.Checkbutton(
            right_frame, text=self.i18n.get('write_metadata'),
            variable=self.write_metadata, style="Modern.TCheckbutton"
        )
        self.widgets['metadata_check'].pack(anchor="w", pady=(5, 0))
    
    def _build_progress_section(self, parent: ttk.Frame) -> None:
        """Build progress and log section / 构建进度和日志区域"""
        progress_frame = ttk.Frame(parent, style="Modern.TFrame")
        progress_frame.pack(fill=tk.X, pady=(15, 10))
        
        self.widgets['progress_label'] = ttk.Label(progress_frame, text=self.i18n.get('ready'), style="Modern.TLabel")
        self.widgets['progress_label'].pack(anchor="w")
        
        self.progress_bar = ttk.Progressbar(
            progress_frame, style="Modern.Horizontal.TProgressbar", mode="determinate", length=400
        )
        self.progress_bar.pack(fill=tk.X, pady=(8, 0))
        
        log_frame = ttk.Frame(parent, style="Modern.TFrame")
        log_frame.pack(fill=tk.BOTH, expand=True, pady=(10, 0))
        
        self.widgets['log_label'] = ttk.Label(log_frame, text=self.i18n.get('log_title'), style="Modern.TLabel")
        self.widgets['log_label'].pack(anchor="w", pady=(0, 5))
        
        log_container = tk.Frame(log_frame, bg=ModernStyle.BORDER)
        log_container.pack(fill=tk.BOTH, expand=True)
        
        self.log_text = scrolledtext.ScrolledText(
            log_container, font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_SMALL),
            bg=ModernStyle.BG_TERTIARY, fg=ModernStyle.TEXT_SECONDARY,
            insertbackground=ModernStyle.TEXT_PRIMARY, relief="flat", height=8, wrap=tk.WORD
        )
        self.log_text.pack(fill=tk.BOTH, expand=True, padx=1, pady=1)
        self.log_text.configure(state="disabled")
    
    def _build_action_buttons(self, parent: ttk.Frame) -> None:
        """Build action buttons / 构建操作按钮"""
        button_frame = ttk.Frame(parent, style="Modern.TFrame")
        button_frame.pack(fill=tk.X, pady=(15, 10))
        
        self.widgets['start_btn'] = tk.Button(
            button_frame, text=self.i18n.get('start'), command=self._start_processing,
            bg=ModernStyle.ACCENT, fg="white",
            activebackground=ModernStyle.ACCENT_HOVER, activeforeground="white",
            relief="flat", cursor="hand2",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_LARGE, "bold"), padx=30, pady=12
        )
        self.widgets['start_btn'].pack(side=tk.RIGHT)
        
        self.widgets['clear_btn'] = tk.Button(
            button_frame, text=self.i18n.get('clear_log'), command=self._clear_log,
            bg=ModernStyle.BG_TERTIARY, fg=ModernStyle.TEXT_SECONDARY,
            activebackground=ModernStyle.BORDER, activeforeground=ModernStyle.TEXT_PRIMARY,
            relief="flat", cursor="hand2",
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_NORMAL), padx=16, pady=10
        )
        self.widgets['clear_btn'].pack(side=tk.LEFT)
    
    def _create_card(self, parent: ttk.Frame, title_key: str) -> Tuple[tk.Frame, tk.Frame]:
        """Create card component / 创建卡片组件"""
        outer = tk.Frame(parent, bg=ModernStyle.BORDER)
        outer.pack(fill=tk.X, pady=(0, 12))
        
        card = tk.Frame(outer, bg=ModernStyle.BG_SECONDARY)
        card.pack(fill=tk.X, padx=1, pady=1)
        
        content = tk.Frame(card, bg=ModernStyle.BG_SECONDARY)
        content.pack(fill=tk.X, padx=20, pady=12)
        
        title_label = tk.Label(
            content, text=self.i18n.get(title_key),
            bg=ModernStyle.BG_SECONDARY, fg=ModernStyle.TEXT_PRIMARY,
            font=(ModernStyle.FONT_FAMILY, ModernStyle.FONT_SIZE_LARGE, "bold")
        )
        title_label.pack(anchor="w", pady=(0, 8))
        self.widgets[title_key] = title_label
        
        return card, content
    
    def _toggle_language(self) -> None:
        """Toggle language / 切换语言"""
        self.i18n.toggle()
        self._update_ui_language()
    
    def _update_ui_language(self) -> None:
        """Update UI language / 更新界面语言"""
        self.root.title(self.i18n.get('app_title'))
        
        # Update widgets / 更新组件
        updates = {
            'title': 'header_title',
            'subtitle': 'header_subtitle',
            'api_card_title': 'api_card_title',
            'api_hint': 'api_hint',
            'dir_card_title': 'dir_card_title',
            'options_card_title': 'options_card_title',
            'mode_label': 'processing_mode',
            'other_label': 'other_options',
            'log_label': 'log_title',
        }
        
        for widget_key, i18n_key in updates.items():
            if widget_key in self.widgets:
                self.widgets[widget_key].configure(text=self.i18n.get(i18n_key))
        
        # Update buttons / 更新按钮
        if 'browse_btn' in self.widgets:
            self.widgets['browse_btn'].configure(text=self.i18n.get('browse'))
        if 'clear_btn' in self.widgets:
            self.widgets['clear_btn'].configure(text=self.i18n.get('clear_log'))
        if 'start_btn' in self.widgets:
            btn_text = self.i18n.get('processing_btn') if self.is_processing else self.i18n.get('start')
            self.widgets['start_btn'].configure(text=btn_text)
        if 'lang_btn' in self.widgets:
            lang_text = "中文" if self.i18n.lang == 'zh' else "English"
            self.widgets['lang_btn'].configure(text=self.i18n.get('language') + ": " + lang_text)
        if 'show_key_btn' in self.widgets:
            key_text = self.i18n.get('hide_key') if self.show_key.get() else self.i18n.get('show_key')
            self.widgets['show_key_btn'].configure(text=key_text)
        
        # Update radio buttons / 更新单选按钮
        if 'batch_radio' in self.widgets:
            self.widgets['batch_radio'].configure(text=self.i18n.get('batch_mode'))
        if 'single_radio' in self.widgets:
            self.widgets['single_radio'].configure(text=self.i18n.get('single_mode'))
        if 'metadata_check' in self.widgets:
            self.widgets['metadata_check'].configure(text=self.i18n.get('write_metadata'))
        
        # Update progress label / 更新进度标签
        if 'progress_label' in self.widgets and not self.is_processing:
            self.widgets['progress_label'].configure(text=self.i18n.get('ready'))
        
        # Update file count / 更新文件计数
        self._update_file_count()
    
    def _toggle_key_visibility(self) -> None:
        """Toggle key visibility / 切换密钥可见性"""
        self.show_key.set(not self.show_key.get())
        self.api_entry.configure(show="" if self.show_key.get() else "•")
        key_text = self.i18n.get('hide_key') if self.show_key.get() else self.i18n.get('show_key')
        self.widgets['show_key_btn'].configure(text=key_text)
    
    def _browse_directory(self) -> None:
        """Browse directory / 浏览选择目录"""
        directory = filedialog.askdirectory(
            title=self.i18n.get('select_dir_hint'),
            initialdir=self.target_dir.get() or os.getcwd()
        )
        if directory:
            self.target_dir.set(directory)
            self._update_file_count()
    
    def _update_file_count(self) -> None:
        """Update file count / 更新文件计数"""
        dir_path = Path(self.target_dir.get())
        if not dir_path.is_dir():
            self.widgets['file_count'].configure(text=self.i18n.get('select_dir_hint'))
            return
        
        count = 0
        for ext in CONFIG.supported_extensions:
            count += len(list(dir_path.glob(f"**/*{ext}")))
        
        if count > 0:
            self.widgets['file_count'].configure(text=self.i18n.get('files_found', count=count))
        else:
            self.widgets['file_count'].configure(text=self.i18n.get('no_files_found'))
    
    def _clear_log(self) -> None:
        """Clear log / 清除日志"""
        self.log_text.configure(state="normal")
        self.log_text.delete(1.0, tk.END)
        self.log_text.configure(state="disabled")
    
    def _log(self, message: str) -> None:
        """Add log message / 添加日志消息"""
        self.log_text.configure(state="normal")
        self.log_text.insert(tk.END, message + "\n")
        self.log_text.see(tk.END)
        self.log_text.configure(state="disabled")
    
    def _update_log(self) -> None:
        """Update log from queue / 从队列更新日志"""
        try:
            while True:
                message = self.log_queue.get_nowait()
                self._log(message)
        except queue.Empty:
            pass
        self.root.after(100, self._update_log)
    
    def _update_progress(self, current: int, total: int, message: str = "") -> None:
        """Update progress / 更新进度"""
        self.progress_bar["maximum"] = total
        self.progress_bar["value"] = current
        
        if message:
            self.widgets['progress_label'].configure(text=message)
        else:
            percent = (current / total * 100) if total > 0 else 0
            self.widgets['progress_label'].configure(
                text=self.i18n.get('progress', current=current, total=total, percent=percent)
            )
    
    def _start_processing(self) -> None:
        """Start processing / 开始处理"""
        if self.is_processing:
            return
        
        if not self.api_keys.get().strip():
            messagebox.showerror(self.i18n.get('error'), self.i18n.get('api_key_required'))
            return
        
        if not self.target_dir.get().strip():
            messagebox.showerror(self.i18n.get('error'), self.i18n.get('dir_required'))
            return
        
        target_path = Path(self.target_dir.get())
        if not target_path.is_dir():
            messagebox.showerror(self.i18n.get('error'), self.i18n.get('dir_not_exist'))
            return
        
        self.is_processing = True
        self.widgets['start_btn'].configure(state="disabled", text=self.i18n.get('processing_btn'))
        self._clear_log()
        
        thread = threading.Thread(target=self._run_processing, daemon=True)
        thread.start()
    
    def _run_processing(self) -> None:
        """Run processing in new thread / 在新线程中运行处理"""
        try:
            asyncio.run(self._async_processing())
        except Exception as e:
            self.logger.error(self.i18n.get('processing_error', error=e))
        finally:
            self.root.after(0, self._processing_complete)
    
    async def _async_processing(self) -> None:
        """Async processing main logic / 异步处理主逻辑"""
        keys = [k.strip() for k in self.api_keys.get().split(',') if k.strip()]
        if not keys:
            self.logger.error(self.i18n.get('no_valid_keys'))
            return
        
        key_manager = APIKeyManager(keys, CONFIG.tracker_file)
        self.logger.info(self.i18n.get('api_keys_found', count=len(keys)))
        
        target_dir = Path(self.target_dir.get())
        mode = ProcessingMode.BATCH if self.mode.get() == "batch" else ProcessingMode.SINGLE
        write_metadata = self.write_metadata.get()
        
        mode_str = self.i18n.get('mode_batch') if mode == ProcessingMode.BATCH else self.i18n.get('mode_single')
        self.logger.info(self.i18n.get('processing_mode_log', mode=mode_str))
        
        metadata_status = self.i18n.get('metadata_enabled') if write_metadata else self.i18n.get('metadata_disabled')
        self.logger.info(self.i18n.get('metadata_log', status=metadata_status))
        
        pending_manager = PendingFilesManager(CONFIG.pending_files_log)
        pending = pending_manager.load()
        
        if pending:
            self.logger.info(self.i18n.get('resume_from_breakpoint', count=len(pending)))
            file_paths = [p for p in pending if p.exists()]
        else:
            file_paths = []
            for ext in CONFIG.supported_extensions:
                file_paths.extend(target_dir.glob(f"**/*{ext}"))
            file_paths = sorted(file_paths, key=str)
        
        if not file_paths:
            self.logger.info(self.i18n.get('no_files_to_process'))
            return
        
        self.logger.info(self.i18n.get('files_to_process', count=len(file_paths)))
        self.root.after(0, lambda: self._update_progress(0, len(file_paths), self.i18n.get('preparing')))
        
        for key in keys:
            if MODEL.configure(key):
                self.logger.info(self.i18n.get('api_configured'))
                break
        else:
            self.logger.error(self.i18n.get('all_keys_invalid'))
            return
        
        self.logger.info(self.i18n.get('extracting_text'))
        loop = asyncio.get_running_loop()
        file_items = []
        
        with ThreadPoolExecutor(max_workers=CONFIG.io_workers) as pool:
            tasks = [loop.run_in_executor(pool, extract_and_count, p, CONFIG) for p in file_paths]
            results = await asyncio.gather(*tasks)
        
        for result in results:
            if result and result.tokens <= CONFIG.max_tokens_per_request:
                file_items.append(result)
        
        self.logger.info(self.i18n.get('extracted_files', count=len(file_items)))
        
        if not file_items:
            self.logger.warning(self.i18n.get('no_content_extracted'))
            return
        
        limiter = RateLimiter(CONFIG.rpm_limit, CONFIG.tpm_limit)
        renamer = FileRenamer(write_metadata)
        processor = FileProcessor(CONFIG, limiter, renamer, self.logger, self.i18n)
        
        remaining: Deque[FileItem] = deque(file_items)
        total_processed = 0
        processed_count = 0
        
        def update_progress(count: int):
            nonlocal processed_count
            processed_count += count
            self.root.after(0, lambda: self._update_progress(
                processed_count, len(file_items), self.i18n.get('processing')
            ))
        
        for key_idx, api_key in enumerate(keys):
            if not remaining:
                break
            
            self.logger.info(self.i18n.get('using_key', idx=key_idx + 1, total=len(keys)))
            
            if not MODEL.configure(api_key):
                continue
            
            quota = key_manager.get_remaining_quota(api_key, CONFIG.daily_request_limit)
            if quota <= 0:
                self.logger.warning(self.i18n.get('quota_exhausted'))
                continue
            
            self.logger.info(self.i18n.get('remaining_quota', quota=quota))
            
            current_items = list(remaining)
            remaining = deque()
            
            if mode == ProcessingMode.BATCH:
                batches = pack_batches_ffd(current_items, CONFIG.max_tokens_per_request, CONFIG.max_items_per_batch)
                batches_to_process = batches[:quota]
                for b in batches[quota:]:
                    remaining.extend(b.items)
                
                for batch in batches_to_process:
                    key_manager.increment_usage(api_key)
                    result = await processor.process_batch(batch, update_progress)
                    
                    if result.success:
                        total_processed += len(batch.items)
                    else:
                        remaining.extend(result.failed_items)
                        if result.quota_exceeded:
                            break
            else:
                for item in current_items:
                    if quota <= 0:
                        remaining.append(item)
                        continue
                    
                    key_manager.increment_usage(api_key)
                    quota -= 1
                    
                    result = await processor.process_single(item, update_progress)
                    
                    if result.success:
                        total_processed += 1
                    else:
                        if result.failed_item:
                            remaining.append(result.failed_item)
                        if result.quota_exceeded:
                            remaining.extend(current_items[current_items.index(item) + 1:])
                            break
            
            key_manager.save_tracker()
        
        if remaining:
            pending_manager.save([item.path for item in remaining])
            self.logger.warning(self.i18n.get('remaining_files', count=len(remaining)))
        else:
            pending_manager.clear()
        
        self.logger.info(self.i18n.get('completed', count=total_processed))
    
    def _processing_complete(self) -> None:
        """Processing complete callback / 处理完成回调"""
        self.is_processing = False
        self.widgets['start_btn'].configure(state="normal", text=self.i18n.get('start'))
        self._update_file_count()
    
    def run(self) -> None:
        """Run application / 运行应用"""
        self.root.update_idletasks()
        width = self.root.winfo_width()
        height = self.root.winfo_height()
        x = (self.root.winfo_screenwidth() // 2) - (width // 2)
        y = (self.root.winfo_screenheight() // 2) - (height // 2)
        self.root.geometry(f"{width}x{height}+{x}+{y}")
        
        self.root.mainloop()


# ============================================================================
# Main Entry / 主入口
# ============================================================================

def main():
    """Main entry function / 主入口函数"""
    app = GeminiRenamerGUI()
    app.run()


if __name__ == "__main__":
    main()

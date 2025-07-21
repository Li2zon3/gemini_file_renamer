# --------------------------------------------------------------------------------
# gemini_file_renamer_op
#
# 功能:
# - 双模式支持: 可选择高效批处理模式 (默认) 和简单的单文件处理模式 (--mode single)。
# - 元数据写入：可选择将提取的信息写入到PDF、DOCX、EPUB/AZW3文件的元数据中。
# - 并发控制: 使用Semaphore精确控制并发任务数，保证程序稳定高效；自动处理API速率限制（RPM/TPM）和每日请求数限制。
# - 支持多个API密钥，主动追踪每个密钥的每日API使用情况（记录在 request_tracker.json），在一个密钥的每日免费额度用尽时自动切换到下一个。
# - 同时提取文件开头和结尾的文本内容，以提高识别准确率；根据gemini模型识别出的信息生成规范化文件名。
# - 支持强大的断点续传（按天、按密钥、按文件列表追踪）和多遍失败重试，当批处理失败时（如API返回结果数不匹配），会将该批次留给下一个API密钥处理。
#
# 使用前请先安装必要的库:
# pip install google-generativeai pymupdf pathvalidate tqdm python-docx EbookLib beautifulsoup4
#
# 运行: 默认批处理模式、元数据写入
# python gemini_renamer_op.py [可选的文件夹路径] [--mode single] [--no-metadata]
# --------------------------------------------------------------------------------

import os
import sys
import json
import time
import asyncio
import logging
import argparse
from pathlib import Path
from datetime import date
from collections import deque
from bs4 import BeautifulSoup

# 第三方库
import google.generativeai as genai
import pymupdf  # PyMuPDF for PDF
from docx import Document  # python-docx for DOCX
from ebooklib import epub, ITEM_DOCUMENT  # EbookLib for EPUB/AZW3
from pathvalidate import sanitize_filename
from tqdm.asyncio import tqdm


# --- 1. 初始化与配置 ---
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def configure_api_keys():
    """从环境变量或用户输入中获取并配置API密钥列表。"""
    keys_str = os.getenv("GOOGLE_API_KEY")
    if not keys_str:
        print("-----------------------------------------------------------------")
        print("未找到 GOOGLE_API_KEY 环境变量。")
        keys_str = input("请输入您的一个或多个 Google API 密钥 (若有多个，请用逗号','分隔):\n").strip()
        print("-----------------------------------------------------------------")

    if not keys_str:
        logging.error("错误：未提供任何 API 密钥，程序即将退出。")
        sys.exit(1)

    api_keys = [key.strip() for key in keys_str.split(',') if key.strip()]

    if not api_keys:
        logging.error("错误：提供的 API 密钥为空，程序即将退出。")
        sys.exit(1)

    logging.info(f"找到 {len(api_keys)} 个 API 密钥。将进行多遍处理。")
    return api_keys

async def switch_and_configure_api(api_key):
    """配置并验证指定的API密钥。"""
    global MODEL
    try:
        genai.configure(api_key=api_key)
        MODEL = genai.GenerativeModel('models/gemini-2.5-flash')
        MODEL.api_key = api_key # 自定义属性，用于日志记录
        logging.info(f"API 密钥 (前8位: {api_key[:8]}...) 配置成功。模型: gemini-2.5-flash")
        return True
    except Exception as e:
        logging.error(f"API 密钥 (前8位: {api_key[:8]}...) 配置失败。错误: {e}")
        return False


# --- 2. 全局常量与模型定义 ---
RPM_LIMIT = 10
TPM_LIMIT = 250000
DAILY_REQUEST_LIMIT = 250
MAX_TOKENS_PER_REQUEST = 28000 # 单个请求（批处理或单文件）的安全Token上限
CONCURRENCY_LIMIT = 10 # 并发任务数
MAX_RETRIES = 3

SUPPORTED_EXTENSIONS = ['.pdf', '.epub', '.azw3', '.docx']
MODEL = None

# --- Prompts and JSON Schemas ---
API_PROMPT_INSTRUCTION_BATCH = """
Analyze the following text, which contains MULTIPLE documents concatenated together.
Each document starts with a "--- START OF FILE: [filename] ---" marker and ends with an "--- END OF FILE: [filename] ---" marker.
For EACH document provided, extract its metadata. Crucially, also extract a list of 3-5 relevant keywords from the document's content.
Return a single JSON array (a list) containing all the extracted JSON objects.
The order of objects in the final list MUST match the order of the documents in the input text.
Do not add any commentary. Only return the JSON array.
"""

API_PROMPT_INSTRUCTION_SINGLE = """
Analyze the text from the following document to extract its metadata.
Based on the content, provide a JSON object with the following details.
Crucially, also extract a list of 3-5 relevant keywords from the document's content.
Do not add any commentary. Only return the JSON object.
"""

SINGLE_OBJECT_SCHEMA = {
    "type": "object",
    "properties": {
        "title": {"type": "string"},
        "authors": {"type": "array", "items": {"type": "string"}},
        "keywords": {"type": "array", "items": {"type": "string"}}, # 新增
        "translators": {"type": "string"},
        "editors": {"type": "string"},
        "publisher_or_journal": {"type": "string"},
        "journal_volume_issue": {"type": "string"},
        "publication_date": {"type": "string"},
        "start_page": {"type": "integer"}
    },
    "required": ["title"]
}

JSON_SCHEMA_BATCH = {
    "type": "array",
    "items": SINGLE_OBJECT_SCHEMA
}

# --- 3. 速率控制器 ---
class RateLimiter:
    """控制对API的请求速率，避免超出限制。"""
    def __init__(self, rpm, tpm):
        self.rpm = rpm
        self.tpm = tpm
        self.request_timestamps = deque()
        self.token_timestamps = deque()

    async def wait_for_slot(self, tokens_needed):
        """
        优化后的版本：精确计算等待时间，而不是0.1秒轮询检查。
        这可以显著降低CPU在等待期间的占用率。
        """
        while True:
            now = time.time()
            # 清理超过一分钟的旧时间戳
            while self.request_timestamps and self.request_timestamps[0] < now - 60:
                self.request_timestamps.popleft()
            while self.token_timestamps and self.token_timestamps[0][0] < now - 60:
                self.token_timestamps.popleft()

            current_requests = len(self.request_timestamps)
            current_tokens = sum(t[1] for t in self.token_timestamps)

            # 如果有足够的配额，则立即获得槽位并退出循环
            if current_requests < self.rpm and (current_tokens + tokens_needed) <= self.tpm:
                self.request_timestamps.append(now)
                self.token_timestamps.append((now, tokens_needed))
                break

            # --- 精确计算等待时间 ---
            rpm_wait = 0
            if current_requests >= self.rpm and self.request_timestamps:
                # 计算需要等待多久，最旧的请求才会离开60秒窗口
                rpm_wait = (self.request_timestamps[0] + 60) - now

            tpm_wait = 0
            if (current_tokens + tokens_needed) > self.tpm and self.token_timestamps:
                # 计算需要释放多少token
                tokens_to_free = (current_tokens + tokens_needed) - self.tpm
                freed_tokens = 0
                wait_until_ts = 0
                # 遍历token记录，找到释放足够token所需的最晚时间点
                for ts, tk in self.token_timestamps:
                    freed_tokens += tk
                    if freed_tokens >= tokens_to_free:
                        wait_until_ts = ts
                        break
                if wait_until_ts > 0:
                    tpm_wait = (wait_until_ts + 60) - now
            
            # 取最长的等待时间，确保RPM和TPM都满足要求
            wait_time = max(0.1, rpm_wait, tpm_wait)
            logging.info(f"速率限制已达上限。将精确等待 {wait_time:.2f} 秒...")
            await asyncio.sleep(wait_time)


# --- 4. 核心异步功能函数 ---
async def process_batch(batch, batch_tokens, limiter, pbar, semaphore, write_metadata_flag):
    """(批处理模式) 异步处理单个批次的文件，并使用Semaphore控制并发。"""
    async with semaphore:
        if not batch or MODEL is None:
            pbar.update(len(batch))
            return {"success": False, "failed_items": batch, "quota_exceeded": False}

        prompt_parts = [API_PROMPT_INSTRUCTION_BATCH]
        for item in batch:
            prompt_parts.append(f"\n\n--- START OF FILE: {item['path'].name} ---\n")
            prompt_parts.append(item['text'])
            prompt_parts.append(f"\n--- END OF FILE: {item['path'].name} ---")
        full_prompt = "".join(prompt_parts)
        generation_config = {"response_mime_type": "application/json", "response_schema": JSON_SCHEMA_BATCH}

        for attempt in range(MAX_RETRIES):
            try:
                await limiter.wait_for_slot(batch_tokens)
                response = await MODEL.generate_content_async(full_prompt, generation_config=generation_config)
                results = json.loads(response.text)

                if not isinstance(results, list) or len(results) != len(batch):
                    logging.warning(f"批处理返回结果数不匹配。预期 {len(batch)} 个，得到 {len(results)} 个。此批次将留待下一轮尝试。")
                    pbar.update(len(batch))
                    return {"success": False, "failed_items": batch, "quota_exceeded": False}

                for i, info in enumerate(results):
                    original_item = batch[i]
                    new_name = build_filename(info)
                    await rename_and_write_metadata(original_item['path'], new_name, info, write_metadata_flag)

                pbar.update(len(batch))
                return {"success": True, "failed_items": [], "quota_exceeded": False}

            except json.JSONDecodeError:
                logging.error(f"批处理JSON解析失败: API返回: {getattr(response, 'text', 'N/A')[:200]}...")
                break
            except Exception as e:
                logging.error(f"处理批次时出错 (尝试 {attempt + 1}/{MAX_RETRIES}): {e}")
                is_quota_error = "quota" in str(e).lower() or "exceeded" in str(e).lower() or "429" in str(e)
                if is_quota_error:
                    logging.warning(f"API密钥 (前8位: {MODEL.api_key[:8]}...) 配额可能已用尽。")
                    return {"success": False, "failed_items": batch, "quota_exceeded": True}
                if attempt < MAX_RETRIES - 1:
                    await asyncio.sleep(2 ** (attempt + 1))
                else:
                    logging.error("已达到最大重试次数，放弃此批次。")
                    break
        
        pbar.update(len(batch))
        return {"success": False, "failed_items": batch, "quota_exceeded": False}

async def process_single_file(file_item, limiter, pbar, write_metadata_flag):
    """(单文件模式) 异步处理单个文件。"""
    if not file_item or MODEL is None:
        pbar.update(1)
        return {"success": False, "failed_item": file_item, "quota_exceeded": False}

    prompt_parts = [API_PROMPT_INSTRUCTION_SINGLE, "\n\n", file_item['text']]
    full_prompt = "".join(prompt_parts)
    generation_config = {"response_mime_type": "application/json", "response_schema": SINGLE_OBJECT_SCHEMA}

    for attempt in range(MAX_RETRIES):
        try:
            await limiter.wait_for_slot(file_item['tokens'])
            response = await MODEL.generate_content_async(full_prompt, generation_config=generation_config)
            info = json.loads(response.text)
            
            new_name = build_filename(info)
            await rename_and_write_metadata(file_item['path'], new_name, info, write_metadata_flag)
            
            pbar.update(1)
            return {"success": True, "failed_item": None, "quota_exceeded": False}

        except json.JSONDecodeError:
            logging.error(f"文件 {file_item['path'].name} 的JSON解析失败: API返回: {getattr(response, 'text', 'N/A')[:200]}...")
            break
        except Exception as e:
            logging.error(f"处理文件 {file_item['path'].name} 时出错 (尝试 {attempt + 1}/{MAX_RETRIES}): {e}")
            is_quota_error = "quota" in str(e).lower() or "exceeded" in str(e).lower() or "429" in str(e)
            if is_quota_error:
                logging.warning(f"API密钥 (前8位: {MODEL.api_key[:8]}...) 配额可能已用尽。")
                return {"success": False, "failed_item": file_item, "quota_exceeded": True}
            if attempt < MAX_RETRIES - 1:
                await asyncio.sleep(2 ** (attempt + 1))
            else:
                logging.error(f"文件 {file_item['path'].name} 已达到最大重试次数，放弃。")
                break
    
    pbar.update(1)
    return {"success": False, "failed_item": file_item, "quota_exceeded": False}


# --- 5. 辅助函数 ---
# --- 5a. 每日请求跟踪模块 ---
PENDING_FILES_LOG = Path("./pending_files.txt")
TRACKER_FILE = Path("./request_tracker.json")

def load_request_tracker():
    today_str = date.today().isoformat()
    default_tracker = {"date": today_str, "usage": {}}
    if not TRACKER_FILE.exists(): return default_tracker
    try:
        with open(TRACKER_FILE, 'r', encoding='utf-8') as f: tracker = json.load(f)
        if tracker.get("date") != today_str:
            logging.info("新的一天，重置所有API密钥的每日请求计数器。")
            return default_tracker
        if "usage" not in tracker: tracker["usage"] = {}
        return tracker
    except (json.JSONDecodeError, IOError) as e:
        logging.warning(f"读取请求跟踪文件失败，将重新开始计数。错误: {e}")
        return default_tracker

def save_request_tracker(tracker_data):
    try:
        with open(TRACKER_FILE, 'w', encoding='utf-8') as f: json.dump(tracker_data, f, indent=4, ensure_ascii=False)
    except IOError as e: logging.error(f"保存请求跟踪文件失败: {e}")

# --- 5b. 文本提取模块 ---
def _extract_from_pdf(pdf_path, num_pages_start=4, num_pages_end=3):
    text_content = []
    try:
        with pymupdf.open(pdf_path) as doc:
            total_pages = doc.page_count
            pages_to_extract = set(range(min(num_pages_start, total_pages)))
            if total_pages > num_pages_start + num_pages_end:
                pages_to_extract.update(range(total_pages - num_pages_end, total_pages))
            
            for i in sorted(list(pages_to_extract)):
                text_content.append(doc[i].get_text(sort=True))
    except Exception as e: logging.error(f"提取PDF文本时出错: {pdf_path.name}, 错误: {e}"); return ""
    return "\n".join(text_content)

def _extract_from_epub(epub_path, num_chapters_start=5, num_chapters_end=4):
    text_content = []
    try:
        book = epub.read_epub(epub_path)
        doc_items = list(book.get_items_of_type(ITEM_DOCUMENT))
        items_to_process = doc_items[:num_chapters_start]
        if len(doc_items) > num_chapters_start + num_chapters_end:
            items_to_process.extend(doc_items[-num_chapters_end:])

        for item in items_to_process:
            soup = BeautifulSoup(item.get_body_content(), 'html.parser')
            text_content.append(soup.get_text("\n", strip=True))
    except Exception as e: logging.error(f"提取 EPUB/AZW3 文本时出错: {epub_path.name}, 错误: {e}"); return ""
    return "\n\n".join(text_content)

def _extract_from_docx(docx_path, num_paras_start=20, num_paras_end=15):
    text_content = []
    try:
        doc = Document(docx_path)
        all_paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        if len(all_paras) > num_paras_start + num_paras_end:
            text_content.extend(all_paras[:num_paras_start])
            text_content.extend(all_paras[-num_paras_end:])
        else:
            text_content = all_paras
    except Exception as e: logging.error(f"提取 DOCX 文本时出错: {docx_path.name}, 错误: {e}"); return ""
    return "\n".join(text_content)

def smart_truncate_text(text: str, max_chars: int) -> str:
    if len(text) <= max_chars:
        return text

    # 计算保留开头60%和结尾40%的字符数
    head_chars = int(max_chars * 0.6)
    tail_chars = int(max_chars * 0.4)
    
    head = text[:head_chars]
    tail = text[-tail_chars:]
    
    return f"{head}\n\n--- 内容已截断 (中间部分省略) ---\n\n{tail}"

def extract_text_from_file(file_path):
    extension = file_path.suffix.lower()
    text_to_extract = ""
    try:
        if extension == '.pdf': text_to_extract = _extract_from_pdf(file_path)
        elif extension in ['.epub', '.azw3']: text_to_extract = _extract_from_epub(file_path)
        elif extension == '.docx': text_to_extract = _extract_from_docx(file_path)
        else: logging.warning(f"不支持的文件类型: {file_path.name}"); return None

        # 假设平均1个token约等于3.5个字符，这是一个保守的估计
        max_chars_for_prompt = int(MAX_TOKENS_PER_REQUEST * 3.5)
        return smart_truncate_text(text_to_extract, max_chars_for_prompt)
    except Exception as e: logging.error(f"提取文本时出错: {file_path.name}, 错误: {e}"); return None

# --- 5c. 文件名构建与重命名模块 ---
JOURNAL_KEYWORDS = [
    # 英文
    "journal", "review", "proceedings", "transactions", "quarterly",
    "annals", "bulletin", "magazine", "advances", "letters", "studies",
    "science", "research", "technology", "medicine", "report", "archives",
    # 中文
    "学报", "法学", "研究", "评论", "科学", "技术", "杂志", "动态",
    "报告", "医学", "经济", "哲学", "历史", "通讯", "汇刊", "纪要"
]

def _get_details_string(info):
    """构建除关键词外的详细信息字符串"""
    details = []
    if v := info.get("publisher_or_journal"): details.append(f"出版/期刊: {v}")
    if v := info.get("journal_volume_issue"): details.append(f"卷期: {v}")
    if v := info.get("publication_date"): details.append(f"日期: {v}")
    if v := info.get("editors"): details.append(f"编者: {v}")
    if v := info.get("translators"): details.append(f"译者: {v}")
    if v := info.get("start_page"): details.append(f"页码: {v}")
    return " | ".join(details)

# ===== DOCX 元数据写入函数 (将在独立线程执行) =====
def write_metadata_to_docx(path, info):
    try:
        doc = Document(path)
        cp = doc.core_properties
        cp.title = info.get('title', '')
        cp.author = "、".join(info.get("authors", []))
        cp.subject = _get_details_string(info)
        cp.keywords = ", ".join(info.get('keywords', []))
        cp.comments = "Metadata updated by Gemini File Renamer"
        doc.save(path)
        logging.info(f"成功写入元数据到 DOCX: '{path.name}'")
    except Exception as e:
        logging.error(f"写入DOCX元数据失败 '{path.name}': {e}")

# ===== EPUB 元数据写入函数 (将在独立线程执行) =====
def write_metadata_to_epub(path, info):
    try:
        book = epub.read_epub(path)
        book.set_title(info.get('title', ''))
        # 清理旧作者并添加新作者
        book.metadata.pop('http://purl.org/dc/elements/1.1/', None)
        for author in info.get("authors", []):
            book.add_author(author)
        
        details_str = _get_details_string(info)
        keywords_str = "Keywords: " + ", ".join(info.get('keywords', []))
        book.add_metadata('DC', 'description', f"{details_str}\n{keywords_str}")
        
        epub.write_epub(path, book)
        logging.info(f"成功写入元数据到 EPUB: '{path.name}'")
    except Exception as e:
        logging.error(f"写入EPUB元数据失败 '{path.name}': {e}")

# ===== PDF元数据写入函数 (将在独立线程执行) =====
def write_metadata_to_pdf(pdf_path, info):
    try:
        with pymupdf.open(pdf_path) as doc:
            metadata = doc.metadata
            authors_str = "、".join(info.get("authors", []))
            
            metadata['title'] = info.get('title', '')
            metadata['author'] = authors_str
            metadata['subject'] = _get_details_string(info) # 详细信息放入主题
            metadata['keywords'] = ", ".join(info.get('keywords', [])) # 单独的关键词
            
            doc.set_metadata(metadata)
            doc.save(doc.name, incremental=True, encryption=pymupdf.PDF_ENCRYPT_KEEP)
        logging.info(f"成功将元数据写入: '{pdf_path.name}'")
    except Exception as e:
        logging.error(f"写入元数据到 '{pdf_path.name}' 时失败: {e}")

# ===== 文件名构建函数  =====
def build_filename(info):
    if not info or not info.get('title'):
        return None
    
    title = info.get("title", "无标题").strip()
    authors_str = "、".join(info.get("authors", []))
    
    main_part = f"{title} - {authors_str}" if authors_str else title
        
    extra_parts = []
    if (t := info.get("translators", "").strip()) and t.lower() != 'null':
        extra_parts.append(f"{t} 译")
    if (e := info.get("editors", "").strip()) and e.lower() != 'null':
        pub = info.get("publisher_or_journal", "").lower()
        if not any(k in pub for k in JOURNAL_KEYWORDS):
            extra_parts.append(f"{e} 编")

    return f"{main_part} ({', '.join(extra_parts)})" if extra_parts else main_part

# ===== 异步元数据写入调度器 =====
async def write_metadata_async(path: Path, info: dict):
    """
    这是一个异步的包装器/调度器。
    它负责将同步的、阻塞的I/O操作(如文件保存)调度到独立的线程池中执行,
    从而不会阻塞主程序的事件循环。
    """
    loop = asyncio.get_running_loop()
    ext = path.suffix.lower()
    
    if ext == '.pdf':
        await loop.run_in_executor(None, write_metadata_to_pdf, path, info)
    elif ext == '.docx':
        await loop.run_in_executor(None, write_metadata_to_docx, path, info)
    elif ext in ['.epub', '.azw3']:
        await loop.run_in_executor(None, write_metadata_to_epub, path, info)

# ===== 主重命名函数，可以调用异步写入 =====
async def rename_and_write_metadata(original_path, new_base_name, full_info, write_metadata_flag):
    if not new_base_name:
        logging.warning(f"无法为 {original_path.name} 构建有效文件名，跳过。")
        return
        
    safe_base_name = sanitize_filename(new_base_name)
    new_path = original_path.with_name(f"{safe_base_name}{original_path.suffix}")
    
    counter = 1
    while new_path.exists() and new_path != original_path:
        new_path = original_path.with_name(f"{safe_base_name}_{counter}{original_path.suffix}")
        counter += 1
        
    did_rename = False
    if new_path != original_path:
        try:
            original_path.rename(new_path)
            logging.info(f"成功重命名: '{original_path.name}' -> '{new_path.name}'")
            did_rename = True
        except OSError as e:
            logging.error(f"重命名文件时出错: {original_path.name} -> {new_path.name}, 错误: {e}")
            return # 如果重命名失败，也无需写入元数据
    else:
        # 文件名无需更改，但可能仍需写入元数据
        did_rename = True 

    # 如果重命名成功（或无需重命名）且元数据写入标志为True
    if did_rename and write_metadata_flag:
        # 调用异步函数来处理元数据写入，不会阻塞后续操作。
        await write_metadata_async(new_path, full_info)


# --- 5d. 断点续传/待处理文件日志 ---
def load_pending_files():
    if not PENDING_FILES_LOG.exists(): return []
    try:
        with open(PENDING_FILES_LOG, 'r', encoding='utf-8') as f:
            return [Path(line.strip()) for line in f if line.strip()]
    except IOError: return []

def save_pending_files(file_paths):
    try:
        with open(PENDING_FILES_LOG, 'w', encoding='utf-8') as f:
            for path in file_paths: f.write(f"{path}\n")
    except IOError as e: logging.error(f"无法写入待处理文件日志: {e}")

def clear_pending_files_log():
    if PENDING_FILES_LOG.exists():
        try: PENDING_FILES_LOG.unlink(); logging.info("待处理文件日志已清空。")
        except OSError as e: logging.error(f"无法清空待处理文件日志: {e}")

def get_args():
    parser = argparse.ArgumentParser(
        description="使用Gemini API批量智能重命名文件并写入元数据。",
        formatter_class=argparse.RawTextHelpFormatter
    )
    parser.add_argument("directory", nargs='?', default="./files_to_rename", 
                        help="包含待处理文件的目录路径 (默认为: ./files_to_rename)")
    parser.add_argument("--mode", choices=['batch', 'single'], default='batch', 
                        help="选择处理模式:\n'batch': 高效批处理模式 (默认)\n'single': 逐个文件处理模式")
    parser.add_argument("--no-metadata", action="store_true",
                        help="禁用元数据写入功能。程序默认会写入元数据。")
    return parser.parse_args()


# --- 6. 主执行逻辑 ---
async def main():
    total_start_time = time.time()
    api_keys = configure_api_keys()
    request_tracker = load_request_tracker()
    args = get_args() 
    target_directory = Path(args.directory)
    write_metadata_enabled = not args.no_metadata
    
    logging.info(f"运行模式: {'批处理 (Batch)' if args.mode == 'batch' else '单文件 (Single)'}")
    if write_metadata_enabled:
        logging.info("元数据写入功能已开启。")
    else:
        logging.info("元数据写入功能已禁用。")

    if not target_directory.is_dir():
        target_directory.mkdir(exist_ok=True)
        logging.info(f"目标目录不存在，已创建: {target_directory}。请将文件放入后重新运行。")
        return

    # 文件扫描与准备
    pending_paths = load_pending_files()
    if pending_paths:
        logging.info(f"检测到断点日志，将只处理上次未完成的 {len(pending_paths)} 个文件。")
        files_to_process_paths = [p for p in pending_paths if p.exists()]
    else:
        logging.info("未检测到断点日志，将扫描整个目录进行新任务。")
        files_to_process_paths = list(set([p for ext in SUPPORTED_EXTENSIONS for p in target_directory.glob(f"**/*{ext}")]))

    if not files_to_process_paths:
        logging.info("待处理文件列表为空，程序结束。")
        if pending_paths: clear_pending_files_log()
        return

    # 文本提取与Token计算
    prep_start_time = time.time()
    print(f"准备处理 {len(files_to_process_paths)} 个文件。开始提取文本并计算Token...")
    first_usable_key_found = False
    for api_key in api_keys:
        if await switch_and_configure_api(api_key):
            first_usable_key_found = True
            break
    if not first_usable_key_found:
        logging.error("所有API密钥均无效，无法计算Token，程序退出。")
        return
    
    all_file_data = []
    # 注意：这里的Token计算是同步的，对于大量文件，也可以考虑异步化
    for file_path in tqdm(files_to_process_paths, desc="提取并计算Token"):
        text = extract_text_from_file(file_path)
        if text:
            try:
                tokens = MODEL.count_tokens(text)
                if tokens.total_tokens > MAX_TOKENS_PER_REQUEST:
                     logging.warning(f"文件 {file_path.name} 的Token数({tokens.total_tokens})过大，已跳过。")
                     continue
                all_file_data.append({'path': file_path, 'text': text, 'tokens': tokens.total_tokens})
            except Exception as e:
                logging.error(f"计算Token时出错 ({file_path.name}): {e}")

    if not all_file_data:
        logging.warning("未能为任何文件成功提取文本和计算Token。")
        return
    
    # 多遍处理核心逻辑
    files_for_next_pass = deque(all_file_data)
    total_files_processed_this_session = 0
    limiter = RateLimiter(RPM_LIMIT, TPM_LIMIT)
    
    prep_time = time.time() - prep_start_time
    api_processing_start_time = time.time()
    
    for key_index, api_key in enumerate(api_keys):
        if not files_for_next_pass:
            logging.info("所有文件已在上几轮处理完毕，提前结束。")
            break
        
        logging.info(f"\n--- 第 {key_index + 1}/{len(api_keys)} 遍处理开始 ---")
        logging.info(f"本轮待处理文件数: {len(files_for_next_pass)}")
        
        if not await switch_and_configure_api(api_key):
            logging.warning(f"密钥 #{key_index + 1} 无效，直接进入下一遍。")
            continue

        processing_queue = files_for_next_pass
        files_for_next_pass = deque()

        requests_made_today = request_tracker["usage"].get(api_key, 0)
        requests_left_today = DAILY_REQUEST_LIMIT - requests_made_today
        if requests_left_today <= 0:
            logging.warning(f"密钥 #{key_index + 1} 今日配额已用尽，本轮所有文件将直接移交下一密钥。")
            files_for_next_pass.extend(processing_queue)
            continue
        
        logging.info(f"此密钥今日剩余请求配额: {requests_left_today} 次。")

        with tqdm(total=len(processing_queue), desc=f"密钥 #{key_index+1} 进度", unit="file") as pbar:
            if args.mode == 'batch':
                semaphore = asyncio.Semaphore(CONCURRENCY_LIMIT)
                while processing_queue and requests_left_today > 0:
                    batch_to_process, batch_tokens = [], 0
                    items_in_batch_count = 0
                    while items_in_batch_count < len(processing_queue):
                        item = processing_queue[items_in_batch_count]
                        if batch_to_process and batch_tokens + item['tokens'] > MAX_TOKENS_PER_REQUEST:
                            break
                        batch_tokens += item['tokens']
                        batch_to_process.append(item)
                        items_in_batch_count += 1
                    
                    if not batch_to_process: break
                    for _ in range(items_in_batch_count): processing_queue.popleft()
                    
                    requests_left_today -= 1
                    requests_made_today += 1
                    
                    res = await process_batch(batch_to_process, batch_tokens, limiter, pbar, semaphore, write_metadata_enabled)

                    if res["success"]:
                        total_files_processed_this_session += len(batch_to_process)
                    else:
                        files_for_next_pass.extend(res["failed_items"])
                        if res["quota_exceeded"]:
                            logging.warning("密钥配额耗尽，本轮处理提前结束。")
                            files_for_next_pass.extend(processing_queue)
                            processing_queue.clear()
            
            else: # single mode
                while processing_queue and requests_left_today > 0:
                    file_item = processing_queue.popleft()
                    requests_left_today -= 1
                    requests_made_today += 1
                    
                    res = await process_single_file(file_item, limiter, pbar, write_metadata_enabled)
                    
                    if res["success"]:
                        total_files_processed_this_session += 1
                    else:
                        files_for_next_pass.append(res["failed_item"])
                        if res["quota_exceeded"]:
                            logging.warning("密钥配额耗尽，本轮处理提前结束。")
                            files_for_next_pass.extend(processing_queue)
                            processing_queue.clear()

        request_tracker["usage"][api_key] = requests_made_today
        save_request_tracker(request_tracker)

    api_processing_time = time.time() - api_processing_start_time

    if files_for_next_pass:
        remaining_files = [item['path'] for item in files_for_next_pass]
        logging.warning(f"所有 {len(api_keys)} 个API密钥均已尝试，最终仍有 {len(remaining_files)} 个文件未处理。")
        logging.warning("这些文件已记录在 pending_files.txt。建议使用 --mode single 模式再次运行以处理它们。")
        save_pending_files(remaining_files)
    else:
        logging.info("所有文件已成功处理，清空待处理文件日志。")
        clear_pending_files_log()
    
    total_run_time = time.time() - total_start_time
    average_rate = total_files_processed_this_session / api_processing_time if api_processing_time > 0 else 0
    
    print("\n-----------------------------------------------------------------")
    print("运行结束！")
    print(f"本次运行共成功处理了 {total_files_processed_this_session} 个文件。")
    if files_for_next_pass: print(f"最终有 {len(files_for_next_pass)} 个文件无法处理，已存入日志。")
    
    print("\n--- 耗时分析 ---")
    print(f"准备阶段 (文本提取、Token计算) 耗时: {prep_time:.2f} 秒")
    print(f"API处理阶段 (网络请求、文件重命名) 耗时: {api_processing_time:.2f} 秒")
    print(f"总运行耗时: {total_run_time:.2f} 秒")
    if total_files_processed_this_session > 0:
        print(f"平均处理速率: {average_rate:.2f} 文件/秒")

if __name__ == "__main__":
    try:
        asyncio.run(main())
    except KeyboardInterrupt:
        print("\n程序被用户中断。建议重新运行以使用断点续传功能。")
        sys.exit(0)
